import pandas as pd
import ast
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_UNDERLINE
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
import regex
import string
from docxcompose.composer import Composer
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
from io import BytesIO
import requests
import hashlib
import time
import random
from validators import url
import os

re_normalize_newlines = re.compile(r"\r\n?")
re_remove_invalid_lines = regex.compile(r'^[^\p{Letter}\p{Number}\[\]\(\)]*$', flags=regex.MULTILINE)
re_compress_newlines = re.compile(r"\n+")

chinese_dun_ordinal = r"[零一二三四五六七八九十百]+、.*"
chinese_is_ordinal = r"[零一二三四五六七八九十百]+是.*"
chinese_bracket_ordinal = r"（[零一二三四五六七八九十百]+）.*"
arabic_dot_ordinal = r"(?:\d+\.)+(?!\d).*"
chinese_dun_ordinal_full_stop = r"[零一二三四五六七八九十百]+、.*?。"
chinese_is_ordinal_full_stop = r"[零一二三四五六七八九十百]+是.*?。"
chinese_bracket_ordinal_full_stop = r"（[零一二三四五六七八九十百]+）.*?。"
arabic_dot_ordinal_full_stop = r"(?:\d+\.)+(?!\d).*?。"

ordinal = f"{chinese_dun_ordinal}|{chinese_is_ordinal}|{chinese_bracket_ordinal}|{arabic_dot_ordinal}"
ordinal_full_stop = f"{chinese_dun_ordinal_full_stop}|{chinese_is_ordinal_full_stop}|{chinese_bracket_ordinal_full_stop}|{arabic_dot_ordinal_full_stop}"

digits_letters_punctuation = r"0-9A-Za-z" + re.escape(string.punctuation)


def convert_to_fullwidth(text):
    """将半角标点符号转换为全角标点符号（不转换数字和字母，且不转换点号'.'）
    
    Args:
        text: 输入文本
    
    Returns:
        str: 转换后的文本
    """
    if not text:
        return text
    
    # 半角到全角的标点符号映射表（不包含数字和字母，不转换空格，不转换'.'）
    halfwidth_to_fullwidth = {
        # 常用标点符号（排除'.'）
        '!': '！', '"': '＂', '#': '＃', '$': '＄', '%': '％',
        '&': '＆', "'": '＇', '(': '（', ')': '）', '*': '＊',
        '+': '＋', ',': '，', '-': '－',
        # 注意：不再转换斜杠'/'，以避免影响路径/公式/分数
        # '/': '／',
        ':': '：', ';': '；', '<': '＜', '=': '＝', '>': '＞',
        '?': '？', '@': '＠', '[': '［', '\\': '＼', ']': '］',
        '^': '＾', '_': '＿', '`': '｀', '{': '｛', '|': '｜',
        '}': '｝', '~': '～'  # 移除空格转换，避免生成全角空格
    }
    
    # 逐字符转换
    result = []
    for char in text:
        # 如果字符在映射表中，使用全角字符；否则保持原字符
        result.append(halfwidth_to_fullwidth.get(char, char))
    
    return ''.join(result)


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    """
    在段落中添加书签
    
    Args:
        paragraph: 目标段落对象
        bookmark_text: 书签显示文本
        bookmark_name: 书签名称（用于链接引用）
    """
    run = paragraph.add_run()
    tag = run._r
    
    # 创建书签开始标记
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.append(start)
    
    # 添加书签文本
    text_element = OxmlElement('w:r')
    text_element.text = bookmark_text
    tag.append(text_element)
    
    # 创建书签结束标记
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    end.set(qn('w:name'), bookmark_name)
    tag.append(end)


def add_internal_hyperlink(paragraph, link_to, text, tooltip=None, font_name="宋体", font_size=12, underline=True, font_color=None):
    """在段落中添加指向文档内部书签的超链接。
    
    Args:
        paragraph: 段落对象
        link_to: 目标书签名称
        text: 显示文本
        tooltip: 鼠标悬停提示
        font_name: 字体名称
        font_size: 字体大小（磅）
        underline: 是否下划线
        font_color: 字体颜色，可选值：'red'/'blue'等、RGB三元组(136,0,0)、HEX字符串('#880000'/'880000')、'rgb(r,g,b)'
    """
    hyperlink = OxmlElement('w:hyperlink')
    
    # 设置锚点属性（指向书签）
    hyperlink.set(qn('w:anchor'), link_to)
    
    # 设置工具提示
    if tooltip:
        hyperlink.set(qn('w:tooltip'), tooltip)
    
    # 创建运行元素（放在超链接内部）
    new_run = OxmlElement('w:r')
    
    # 创建运行属性
    rPr = OxmlElement('w:rPr')
    
    # 设置字体
    font_element = OxmlElement('w:rFonts')
    font_element.set(qn('w:ascii'), font_name)
    font_element.set(qn('w:eastAsia'), font_name)
    font_element.set(qn('w:hAnsi'), font_name)
    rPr.append(font_element)
    
    # 设置字体大小（Word 中 w:sz 为磅值的两倍）
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size * 2))
    rPr.append(szCs)
    
    # 设置下划线
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    # 设置字体颜色
    color = OxmlElement('w:color')
    if font_color:
        color_map = {
            'red': 'FF0000',
            'blue': '0000FF',
            'black': '000000',
            'green': '008000',
            'purple': '800080'
        }
        hex_val = None
        # 颜色名称
        if isinstance(font_color, str) and font_color.lower() in color_map:
            hex_val = color_map[font_color.lower()]
        # #RRGGBB 或 RRGGBB
        elif isinstance(font_color, str) and re.fullmatch(r"#?[0-9A-Fa-f]{6}", font_color.strip()):
            hex_val = font_color.strip().lstrip('#').upper()
        # rgb(r,g,b)
        elif isinstance(font_color, str) and font_color.strip().lower().startswith('rgb'):
            nums = re.findall(r"\d+", font_color)
            if len(nums) == 3:
                r_val, g_val, b_val = [min(255, max(0, int(n))) for n in nums]
                hex_val = f"{r_val:02X}{g_val:02X}{b_val:02X}"
        # (r,g,b) 或 [r,g,b]
        elif isinstance(font_color, (tuple, list)) and len(font_color) == 3:
            r_val, g_val, b_val = [min(255, max(0, int(n))) for n in font_color]
            hex_val = f"{r_val:02X}{g_val:02X}{b_val:02X}"
        
        if hex_val:
            color.set(qn('w:val'), hex_val)
        else:
            color.set(qn('w:themeColor'), 'hyperlink')  # 默认超链接颜色
    else:
        color.set(qn('w:themeColor'), 'hyperlink')  # 默认超链接颜色
    rPr.append(color)
    
    # 将运行属性添加到运行
    new_run.append(rPr)
    
    # 创建并添加文本节点 w:t，确保文本可被 python-docx 识别
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)
    
    # 将运行放入超链接中
    hyperlink.append(new_run)
    
    # 关键修复：将超链接作为段落的直接子元素插入
    paragraph._element.append(hyperlink)


def add_hyperlinks_post_processing(doc, target_bookmark="目录", return_link_config=None):
    """
    后处理阶段：为文档中包含返回链接文本的段落添加内部超链接
    这个函数会在所有内容都创建完成后执行，查找返回链接文本并添加超链接指向指定的目标书签
    
    Args:
        doc: Document对象，已包含所有内容的Word文档
        target_bookmark: 目标书签名称，默认为"目录"
        return_link_config: 返回链接配置，包含文本和格式信息
    """
    try:
        # 设置默认返回链接配置
        if return_link_config is None:
            return_link_config = {
                'text': '返回目录',
                'font_name': '宋体',
                'font_size': 12,
                'alignment': 'right',
                'underline': True
                # 不指定 font_color，使用 Word 主题默认超链接颜色
            }
        
        # 获取返回链接文本
        return_text = return_link_config.get('text', '返回目录')
        font_name = return_link_config.get('font_name', '宋体')
        font_size = return_link_config.get('font_size', 12)
        alignment = return_link_config.get('alignment', 'right')
        underline = return_link_config.get('underline', True)
        font_color = return_link_config.get('font_color', None)  # 获取字体颜色配置
        # 当未提供字体颜色时，保持 None，让 add_internal_hyperlink 使用主题默认颜色
        
        # 遍历文档中的所有段落
        for paragraph in doc.paragraphs:
            # 检查段落是否包含返回链接文本
            if paragraph.text.strip() == return_text:
                # 清除现有内容
                paragraph.clear()
                
                # 重新设置段落格式
                paragraph.style = "Normal"
                
                # 设置段落对齐方式
                if alignment == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(12)
                
                # 添加内部超链接，指向指定的目标书签
                add_internal_hyperlink(
                    paragraph,
                    target_bookmark,  # 使用传入的目标书签
                    return_text,  # 使用配置的返回文本
                    tooltip="点击返回文档目录",
                    font_name=font_name,
                    font_size=font_size,
                    underline=underline,
                    font_color=font_color  # 传递字体颜色配置
                )
                    
        print("✓ 内部超链接后处理完成")
        
    except Exception as e:
        print(f"✗ 内部超链接后处理失败: {e}")


def process_lines(content):
    """处理内容行，提取和格式化文本
    
    Args:
        content: 内容字典，值可能是字符串、float或其他类型
    
    Returns:
        list: 处理后的文本行列表
    """
    lines = []
    try:
        # 首先合并所有非空内容到一个完整的文本中
        full_text = ""
        content_parts = []
        
        for key, value in content.items():
            # 处理非字符串类型的值
            if value is None:
                continue
            
            # 安全地检查pandas NA值
            try:
                if pd.isna(value):
                    continue
            except (ValueError, TypeError):
                # 如果pd.isna()失败（比如对于数组），继续处理
                pass
            
            # 如果是数字类型，转换为字符串
            if isinstance(value, (int, float)):
                value = str(value)
            
            # 确保value是字符串类型
            if not isinstance(value, str):
                value = str(value)
            
            # 删除股票价格三元组（价格/涨跌/涨跌幅），兼容全角/半角括号、逗号、百分号与中英文减号
            stock_price_pattern = r'[（(]\s*[0-9０-９]+(?:[.,．][0-9０-９]+)?\s*[，,]\s*[+\-＋－]?\s*[0-9０-９]+(?:[.,．][0-9０-９]+)?\s*[，,]\s*[+\-＋－]?\s*[0-9０-９]+(?:[.,．][0-9０-９]+)?\s*[％%]?\s*[)）]'
            value = re.sub(stock_price_pattern, '', value)
            
            # 删除股票代码信息，兼容（002594.SZ）、(3931.HK) 等全/半角括号与中点
            stock_code_pattern = r'[（(]\s*\d{4,6}\s*[．\.]\s*[A-Za-z]{1,5}\s*[)）]'
            value = re.sub(stock_code_pattern, '', value)
            
            # 处理各种空格字符，删除全角空格"　"
            value = value.replace('　', '')
            value = re.sub(r'[ \t]+', ' ', value)
            value = '\n'.join(line.strip() for line in value.split('\n'))
            value = re.sub(r'\n\s*\n\s*\n\s*\n+', '\n\n', value)
            
            # 提取http链接并单独成行
            http_links = []
            # 匹配http和https链接的正则表达式
            url_pattern = r'https?://[^\s\u4e00-\u9fff]+'
            
            # 查找所有链接
            links = re.findall(url_pattern, value)
            if links:
                # 将链接从原文本中移除
                for link in links:
                    value = value.replace(link, '')
                    http_links.append(link)
            
            # 统一字符为全角格式（不处理HTTP链接）
            value = convert_to_fullwidth(value)
            
            if value.strip():
                content_parts.append(value.strip())
            
            # 将http链接作为独立的内容部分添加（不与其他内容合并）
            content_parts.extend(http_links)
        
        # 分离HTTP链接和文本内容
        http_links = []
        text_parts = []
        
        for part in content_parts:
            # 检查是否是HTTP链接
            if re.match(r'https?://', part.strip()):
                # 直接添加HTTP链接，不进行进一步处理
                http_links.append(part.strip())
            else:
                # 收集文本内容用于段落处理
                if part.strip():
                    text_parts.append(part.strip())
        
        # 合并文本内容进行段落处理
        full_text = '\n'.join(text_parts)
        
        if not full_text.strip():
            return []
        
        # 按段落分割 - 使用更灵活的分段策略
        # 1. 首先尝试按序号标题分段
        paragraphs = []
        
        # 检查是否包含序号标题
        ordinal_patterns = [
            r'[一二三四五六七八九十]、',  # 单个中文数字+顿号
            r'\d+\.',  # 阿拉伯数字序号
            r'（[一二三四五六七八九十]+）',  # 括号中文序号
        ]
        
        has_ordinals = False
        for pattern in ordinal_patterns:
            if re.search(pattern, full_text):
                has_ordinals = True
                break
        
        if has_ordinals:
            # 按序号分段
            # 分割点：序号标题前的位置
            split_pattern = r'(?=(?:[一二三四五六七八九十]、|\d+\.|（[一二三四五六七八九十]+）))'
            parts = re.split(split_pattern, full_text)
            
            current_paragraph = ""
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                
                # 检查是否是序号开头
                is_ordinal_start = False
                for pattern in ordinal_patterns:
                    if re.match(pattern, part):
                        is_ordinal_start = True
                        break
                
                if is_ordinal_start and current_paragraph:
                    # 保存前一个段落，开始新段落
                    paragraphs.append(current_paragraph.strip())
                    current_paragraph = part
                else:
                    # 继续当前段落，但避免过长
                    if current_paragraph and len(current_paragraph) > 200 and part.strip():
                        # 当前段落已经较长，保存并开始新段落
                        paragraphs.append(current_paragraph.strip())
                        current_paragraph = part
                    else:
                        # 继续当前段落
                        if current_paragraph:
                            current_paragraph += '\n' + part
                        else:
                            current_paragraph = part
            
            # 添加最后一个段落
            if current_paragraph.strip():
                paragraphs.append(current_paragraph.strip())
        else:
            # 没有序号，按双换行符或长度分段
            if '\n\n' in full_text:
                paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
            else:
                # 优先按单换行符分段，然后按句号细分
                if '\n' in full_text:
                    # 按换行符分割，但合并过短的段落
                    line_parts = [p.strip() for p in full_text.split('\n') if p.strip()]
                    current_paragraph = ""
                    
                    for line in line_parts:
                        # 检查是否是新段落的开始（标题性质或特定格式）
                        is_new_paragraph = (
                            re.match(r'^[一二三四五六七八九十]+、', line) or  # 中文序号
                            re.match(r'^\d+[、.]', line) or  # 数字序号
                            re.match(r'^\([一二三四五六七八九十]+\)', line) or  # 括号中文序号
                            re.match(r'^\(\d+\)', line) or  # 括号数字序号
                            len(line) < 30 and ('：' in line or '：' in line)  # 短标题行
                        )
                        
                        if is_new_paragraph and current_paragraph:
                            # 保存当前段落，开始新段落
                            paragraphs.append(current_paragraph.strip())
                            current_paragraph = line
                        elif current_paragraph and len(current_paragraph) > 200:
                            # 当前段落过长，开始新段落
                            paragraphs.append(current_paragraph.strip())
                            current_paragraph = line
                        else:
                            # 继续当前段落
                            if current_paragraph:
                                current_paragraph += '\n' + line
                            else:
                                current_paragraph = line
                    
                    # 添加最后一个段落
                    if current_paragraph.strip():
                        paragraphs.append(current_paragraph.strip())
                else:
                    # 完全没有换行符，按句号分段
                    sentences = re.split(r'([。！？])', full_text)
                    current_paragraph = ""
                    
                    for i in range(0, len(sentences), 2):
                        sentence = sentences[i] if i < len(sentences) else ""
                        punctuation = sentences[i+1] if i+1 < len(sentences) else ""
                        full_sentence = sentence + punctuation
                        
                        # 检查是否应该开始新段落
                        if current_paragraph and len(current_paragraph) > 120 and full_sentence.strip():
                            # 当前段落已经足够长，开始新段落
                            paragraphs.append(current_paragraph.strip())
                            current_paragraph = full_sentence
                        else:
                            current_paragraph += full_sentence
                    
                    # 添加最后一个段落
                    if current_paragraph.strip():
                        paragraphs.append(current_paragraph.strip())
                
        
        # 对所有分段路径的结果进行最终的长度检查和强制分割
        final_paragraphs = []
        for para in paragraphs:
            if len(para) > 300:  # 超过300字符的段落需要进一步分割
                # 按换行符分割
                sub_parts = [p.strip() for p in para.split('\n') if p.strip()]
                if len(sub_parts) > 1:
                    # 有换行符，按换行符分割后再检查长度
                    for sub_part in sub_parts:
                        if len(sub_part) > 200:
                            # 子部分仍然过长，强制分割
                            temp_para = sub_part
                            while len(temp_para) > 200:
                                split_pos = 200
                                for punct in ['。', '！', '？', '，', '；']:
                                    pos = temp_para.rfind(punct, 150, 200)
                                    if pos > 0:
                                        split_pos = pos + 1
                                        break
                                
                                final_paragraphs.append(temp_para[:split_pos].strip())
                                temp_para = temp_para[split_pos:].strip()
                            
                            if temp_para:
                                final_paragraphs.append(temp_para)
                        else:
                            final_paragraphs.append(sub_part)
                else:
                    # 如果没有换行符，按长度强制分割
                    temp_para = para
                    while len(temp_para) > 200:
                        # 找到合适的分割点（句号、逗号等）
                        split_pos = 200
                        for punct in ['。', '！', '？', '，', '；']:
                            pos = temp_para.rfind(punct, 150, 200)
                            if pos > 0:
                                split_pos = pos + 1
                                break
                        
                        final_paragraphs.append(temp_para[:split_pos].strip())
                        temp_para = temp_para[split_pos:].strip()
                    
                    if temp_para:
                        final_paragraphs.append(temp_para)
            else:
                final_paragraphs.append(para)
        
        paragraphs = final_paragraphs
        
        # 如果还是没有合理的分段，就作为一个整体段落
        if not paragraphs:
            paragraphs = [full_text.strip()]
        
        # 合并段落和HTTP链接返回
        result = paragraphs + http_links
        return result
        
    except Exception as e:
         print(f"Error processing value in process_lines: {e}")
         print(f"Content type: {type(content)}, Content keys: {list(content.keys()) if hasattr(content, 'keys') else 'No keys'}")
         import traceback
         traceback.print_exc()
         return []
    
    return lines


def process_all_text_paragraphs(doc, *functions):
    paragraphs = doc.paragraphs
    for paragraph in paragraphs[6:]:
        if any(run.text.strip() for run in paragraph.runs):
            for function in functions:
                function(paragraph)


def copy_run_style(run, new_run):
    new_run.bold = run.bold
    new_run.font.name = run.font.name
    new_run.font.size = run.font.size
    new_run.font.underline = run.font.underline  # 保留下划线格式


def replace_halfwidth_quotes_with_fullwidth(paragraph):
    pattern = re.compile(r'"')
    opening_quote = True
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            text_chunks = pattern.split(text)
            quotes = pattern.findall(text)
            for i, text_chunk in enumerate(text_chunks):
                if text_chunk:
                    new_runs.append((text_chunk, run))
                if i < len(quotes):
                    if opening_quote:
                        new_runs.append(("“", run))
                    else:
                        new_runs.append(("”", run))
                    opening_quote = not opening_quote
    for run in paragraph.runs:
        run.text = ""
    for text_chunk, run in new_runs:
        new_run = paragraph.add_run(text_chunk)
        copy_run_style(run, new_run)


def remove_special_symbols(paragraph):
    pattern = regex.compile(r"[^\p{Letter}\p{Number}\p{Han}\p{Punctuation}\p{Math_Symbol}\p{Currency_Symbol}\p{Z}]")
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            cleaned_text = pattern.sub("", text)
            new_runs.append((cleaned_text, run))
    for run in paragraph.runs:
        run.text = ""
    for cleaned_text, run in new_runs:
        new_run = paragraph.add_run(cleaned_text)
        copy_run_style(run, new_run)


def change_digits_letters_punctuation_to_times_new_roman(paragraph):
    pattern = re.compile(r"([" + digits_letters_punctuation + r"]+)")
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            text_chunks = pattern.split(text)
            for i, text_chunk in enumerate(text_chunks):
                if text_chunk:
                    to_change = (i % 2 == 1)
                    new_runs.append((text_chunk, run, to_change))
    for run in paragraph.runs:
        run.text = ""
    for text_chunk, run, to_change in new_runs:
        new_run = paragraph.add_run(text_chunk)
        copy_run_style(run, new_run)
        if to_change:
            new_run.font.name = "Times New Roman"


def remove_space_between_chinese_and_digits_letters_punctuation(paragraph):
    pattern1 = regex.compile(r"([\p{Han}])\s+([" + digits_letters_punctuation + r"]) ")
    pattern2 = regex.compile(r"([" + digits_letters_punctuation + r"])\s+([\p{Han}])")
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if text:
            text = pattern1.sub(r"\1\2", text)
            text = pattern2.sub(r"\1\2", text)
            new_runs.append((text, run))
    for run in paragraph.runs:
        run.text = ""
    for text, run in new_runs:
        new_run = paragraph.add_run(text)
        copy_run_style(run, new_run)


def normalize_spaces_and_convert_punct_except_period(paragraph):
    """统一清理空格（移除全角空格、合并半角空格）并将除'.'外的半角标点转为全角。
    
    注意：为避免破坏URL，本函数会跳过 http/https 链接子串的转换，仅处理非URL片段。
    """
    url_pattern = re.compile(r'https?://[^\s\u4e00-\u9fff]+')
    new_runs = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        # 清理空白：移除全角空格，合并半角空格/制表为单空格
        text = text.replace('　', '')
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 对非URL片段做标点全角化（保留'.'为半角）
        parts = []
        last = 0
        for m in url_pattern.finditer(text):
            non_url_seg = text[last:m.start()]
            if non_url_seg:
                parts.append(convert_to_fullwidth(non_url_seg))
            parts.append(m.group(0))  # 原样保留URL
            last = m.end()
        tail = text[last:]
        if tail:
            parts.append(convert_to_fullwidth(tail))
        new_text = ''.join(parts) if parts else text
        new_runs.append((new_text, run))
    # 清空原runs并重建
    for run in paragraph.runs:
        run.text = ""
    for text, run in new_runs:
        new_run = paragraph.add_run(text)
        copy_run_style(run, new_run)


def center_image_description_paragraphs(doc):
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs):
        if any(run.text.strip() == "" and run.element.xpath(".//w:drawing") for run in paragraph.runs):
            prev_text_paragraph = None
            j = i - 1
            while j >= 0:
                if paragraphs[j].text.strip():
                    prev_text_paragraph = paragraphs[j]
                    break
                j -= 1

            next_text_paragraph = None
            j = i + 1
            while j < len(paragraphs):
                if paragraphs[j].text.strip():
                    next_text_paragraph = paragraphs[j]
                    break
                j += 1

            if prev_text_paragraph and "。" not in prev_text_paragraph.text:
                prev_text_paragraph.paragraph_format.first_line_indent = Pt(0)
                prev_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if next_text_paragraph and "。" not in next_text_paragraph.text:
                next_text_paragraph.paragraph_format.first_line_indent = Pt(0)
                next_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_to_word(info):
    doc = Document("ab_response_formats/guoziwei_start.docx")
    written_heading_1 = set()
    
    # 模板中已包含目录书签，无需重复添加

    for item in info:
        try:
            # 安全获取heading_1
            try:
                heading_1 = item["heading_1"] if "heading_1" in item and pd.notna(item["heading_1"]) else None
            except (ValueError, TypeError):
                heading_1 = item["heading_1"] if "heading_1" in item and item["heading_1"] is not None else None
            
            heading_2 = item["heading_2"]
            
            # 安全获取title
            try:
                title = item["title"] if "title" in item and pd.notna(item["title"]) else None
            except (ValueError, TypeError):
                title = item["title"] if "title" in item and item["title"] is not None else None
            
            # 安全获取source
            try:
                source = item["source"] if "source" in item and pd.notna(item["source"]) else ""
            except (ValueError, TypeError):
                source = item["source"] if "source" in item and item["source"] is not None else ""
            
            # 安全获取date
            try:
                date = item["date"] if "date" in item and pd.notna(item["date"]) else ""
            except (ValueError, TypeError):
                date = item["date"] if "date" in item and item["date"] is not None else ""
            content = item["content"]
            
            # 跳过完全空白的行（所有关键字段都为空）
            if not heading_2 and not title and not content:
                continue
                
            # 处理content内容（如果存在）
            try:
                has_content = content and pd.notna(content)
            except (ValueError, TypeError):
                has_content = content and content is not None
            
            if has_content:
                processed_content = parse_content(content)
                # 如果处理后的内容为空或无效，设置为空列表
                if not processed_content or not isinstance(processed_content, list):
                    print(f"Warning: Invalid content type {type(processed_content)}, setting to empty")
                    content = []
                else:
                    content = processed_content
            else:
                content = []  # 没有content的行设置为空列表

            # 添加一级标题（如果存在且未写入过）
            if heading_1 and heading_1 not in written_heading_1:
                written_heading_1.add(heading_1)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(heading_1)
                paragraph.style = doc.styles["Heading 1"]
                run.font.name = "楷体"
                run.font.size = Pt(22)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加二级标题（heading_2）
            try:
                has_heading_2 = heading_2 and pd.notna(heading_2)
            except (ValueError, TypeError):
                has_heading_2 = heading_2 and heading_2 is not None
            
            if has_heading_2:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(heading_2)
                paragraph.style = doc.styles["Heading 2"]
                run.font.name = "楷体"
                run.font.size = Pt(15)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(6)

            # 添加三级标题（title对应heading_3）
            if title:
                paragraph = doc.add_paragraph()
                # 格式化标题为【标题】格式
                formatted_title = f"【{title}】"
                run = paragraph.add_run(formatted_title)
                # 使用Word内置的Heading 3样式
                paragraph.style = doc.styles["Heading 3"]
                run.font.name = "宋体"
                run.font.size = Pt(12)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Pt(24)
                paragraph.paragraph_format.line_spacing = 1.25

            # 添加正文内容
            for value in content:
                if value.startswith("temp-images"):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(value, width=Inches(5.0))
                    paragraph.alignment = 1
                else:
                    paragraph = doc.add_paragraph()
                    paragraph.style = "Normal"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.first_line_indent = Pt(24)
                    paragraph.paragraph_format.line_spacing = 1.25

                    if "*" in value:
                        text_chunks = re.split(r"(\*\*.*?\*\*)", value)
                        for text_chunk in text_chunks:
                            if text_chunk.startswith("**") and text_chunk.endswith("**"):
                                text_chunk = text_chunk[2:-2]
                                run = paragraph.add_run(text_chunk)
                                run.bold = True
                            else:
                                text_chunk = text_chunk.replace("*", "")
                                run = paragraph.add_run(text_chunk)
                    else:
                        run = paragraph.add_run(value)
                    run.font.name = "宋体"
                    run.font.size = Pt(12)
            
            # 在正文结尾添加日期（如果存在）
            if date:
                paragraph = doc.add_paragraph()
                paragraph.style = "Normal"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Pt(24)
                paragraph.paragraph_format.line_spacing = 1.25
                run = paragraph.add_run(date)
                run.font.name = "宋体"
                run.font.size = Pt(12)
                paragraph.paragraph_format.space_after = Pt(12)
            
            # 修改：仅当 title 与 content 均非空时才添加“返回目录”占位符
            has_title_for_return = bool(title)
            has_content_for_return = isinstance(content, list) and len(content) > 0
            if has_title_for_return and has_content_for_return:
                # 添加"返回目录"占位段落（后续会添加超链接）
                return_paragraph = doc.add_paragraph()
                return_paragraph.style = "Normal"
                return_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
                return_paragraph.paragraph_format.line_spacing = 1.0  # 单倍行距
                return_paragraph.paragraph_format.space_after = Pt(12)
                
                # 添加占位文本，后处理时会查找此文本并添加超链接
                return_run = return_paragraph.add_run("返回目录")
                return_run.font.name = "宋体"
                return_run.font.size = Pt(12)
                return_run.font.underline = WD_UNDERLINE.SINGLE
        except Exception as e:
            print(f"Error processing value: {e}")

    process_all_text_paragraphs(doc, replace_halfwidth_quotes_with_fullwidth, normalize_spaces_and_convert_punct_except_period, remove_special_symbols, change_digits_letters_punctuation_to_times_new_roman, remove_space_between_chinese_and_digits_letters_punctuation)
    center_image_description_paragraphs(doc)
    
    # 后处理阶段：添加所有内部超链接
    add_hyperlinks_post_processing(doc)
    
    os.makedirs("temp-data", exist_ok=True)
    doc_path = f"temp-data/{now_in_filename()}.docx"
    doc.save(doc_path)
    return doc_path

def append_company_info_and_disclaimer(doc_path):
    doc = Document(doc_path)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    composer = Composer(doc)
    doc_to_append = Document("ab_response_formats/guoziwei_end.docx")
    composer.append(doc_to_append)
    composer.save(doc_path)


def tidy(text):
    """清理和标准化文本内容
    
    Args:
        text: 输入的文本内容，可能是字符串、float或None
    
    Returns:
        str: 清理后的文本字符串
    """
    # 处理非字符串类型的输入
    if text is None:
        return ""
    
    # 安全地检查pandas NA值
    try:
        if pd.isna(text):
            return ""
    except (ValueError, TypeError):
        # 如果pd.isna()失败（比如对于数组），继续处理
        pass
    
    # 如果是数字类型，转换为字符串
    if isinstance(text, (int, float)):
        text = str(text)
    
    # 确保text是字符串类型
    if not isinstance(text, str):
        text = str(text)
    
    text = re_normalize_newlines.sub("\n", text)
    text = re_remove_invalid_lines.sub("", text)
    text = re_compress_newlines.sub("\n", text)
    return text.strip()

def get_lines(content):
    """从内容中提取行并去重
    
    Args:
        content: 输入内容，可能是字符串、float或其他类型
    
    Returns:
        dict: 行号到行内容的字典映射
    """
    # 处理非字符串类型的输入
    if content is None:
        return {}
    
    # 安全地检查pandas NA值
    try:
        if pd.isna(content):
            return {}
    except (ValueError, TypeError):
        # 如果pd.isna()失败（比如对于数组），继续处理
        pass
    
    # 如果是数字类型，转换为字符串
    if isinstance(content, (int, float)):
        content = str(content)
    
    # 确保content是字符串类型
    if not isinstance(content, str):
        content = str(content)
    
    try:
        lines = list(dict.fromkeys(line for line in (line.strip() for line in content.splitlines()) if line))
        return dict(enumerate(lines, 1))
    except Exception as e:
        print(f"Error in get_lines: {e}")
        return {}

def now_in_filename():
    return f"{int(time.time())} {''.join([str(random.randint(0, 9)) for _ in range(6)])}"

def get_images_and_insert_paths(content):
    """处理内容中的图片链接，下载图片并返回本地路径
    
    Args:
        content: 内容字典，键为行号，值为内容
    
    Returns:
        list: 处理后的内容列表，图片链接被替换为本地路径
    """
    image_hashes = set()
    result_list = []
    
    for key in sorted(content.keys()):
        value = content[key]
        
        # 处理pandas数据类型和其他非字符串类型
        if value is None:
            continue
        
        # 安全地检查pandas NA值
        try:
            if pd.isna(value):
                continue
        except (ValueError, TypeError):
            # 如果pd.isna()失败（比如对于数组），继续处理
            pass
            
        # 如果是pandas Series，取第一个值
        if hasattr(value, 'iloc'):
            value = value.iloc[0] if len(value) > 0 else ""
        
        # 确保value是字符串类型
        if not isinstance(value, str):
            value = str(value)
        
        # 检查是否为URL（可能是图片链接）
        try:
            # 使用更安全的URL检查方法
            is_url = value.startswith(('http://', 'https://')) and '.' in value
        except Exception as e:
            result_list.append(str(value))
            continue
            
        if is_url:
            try:
                # 尝试下载并处理为图片
                response = requests.get(value, timeout=10)
                image = Image.open(BytesIO(response.content))
                
                # 过滤太小的图片
                if max(image.size) < 100:
                    continue
                    
                # 调整过大图片的尺寸
                if min(image.size) > 1024:
                    ratio = 1024 / min(image.size)
                    image = image.resize((int(image.size[0] * ratio), int(image.size[1] * ratio)), Image.Resampling.LANCZOS)
                
                # 确定图片格式
                image_format = image.format if image.format in ["JPEG", "PNG"] else "JPEG"
                if image_format == "JPEG" and image.mode != "RGB":
                    image = image.convert("RGB")
                
                # 创建保存目录
                os.makedirs("temp-images", exist_ok=True)
                image_path = f"temp-images/{now_in_filename()}.{image_format.lower()}"
                
                # 检查图片是否重复
                buffer = BytesIO()
                image.save(buffer, format=image_format)
                image_data = buffer.getvalue()
                image_hash = hashlib.md5(image_data).hexdigest()
                
                if image_hash not in image_hashes:
                    image_hashes.add(image_hash)
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    result_list.append(image_path)
                # 如果图片重复，跳过不添加
                
            except Exception as e:
                # 如果不是图片或下载失败，保持原始URL
                print(f"处理URL时出错 {value}: {e}")
                result_list.append(value)
        else:
            # 非URL内容直接添加
            result_list.append(value)
    
    return result_list

def parse_content(content):
    """解析内容，处理文本分段和图片链接
    
    Args:
        content: 原始内容，可能是字符串、pandas Series或其他类型
    
    Returns:
        list: 处理后的内容列表，包含文本段落和图片路径
    """
    # 处理pandas数据类型和其他非字符串类型
    if content is None:
        return []
    
    # 安全地检查pandas NA值
    try:
        if pd.isna(content):
            return []
    except (ValueError, TypeError):
        # 如果pd.isna()失败（比如对于数组），继续处理
        pass
    
    # 如果是pandas Series，取第一个值
    if hasattr(content, 'iloc'):
        content = content.iloc[0] if len(content) > 0 else ""
    
    # 确保content是字符串类型
    if not isinstance(content, str):
        content = str(content)
    
    # 首先使用process_lines处理文本内容（包括提取http链接）
    processed_lines = process_lines({"content": content})
    
    # 然后使用get_images_and_insert_paths处理图片链接
    # 将列表转换为字典格式以兼容get_images_and_insert_paths
    lines_dict = {i: line for i, line in enumerate(processed_lines)}
    
    # 处理图片链接并返回最终结果
    return get_images_and_insert_paths(lines_dict)

def manage_thread(requests, thread_count=20):
    if requests:
        with ThreadPoolExecutor(min(len(requests), thread_count)) as executor:
            futures = [(executor.submit(function, *arguments), function, arguments) for function, *arguments in requests]
            return [(future.result(), function.__name__, arguments) for future, function, arguments in futures]
    return []

def parse_contents(contents):
    requests = [(parse_content, content) for content in (contents if isinstance(contents, list) else [contents])]
    return [result for result, name, arguments in manage_thread(requests)]

def extract_from_csv(csv_path):
    """
    从CSV文件中提取数据并返回结构化的记录列表
    
    Args:
        csv_path (str): CSV文件路径
        
    Returns:
        list: 包含字典的列表，每个字典代表一条记录
    """
    # 手动读取CSV文件以确保读取所有行
    import csv
    
    data_rows = []
    try:
        with open(csv_path, 'r', encoding='utf-8', newline='') as file:
            csv_reader = csv.DictReader(file)
            for row in csv_reader:
                data_rows.append(row)
    except UnicodeDecodeError:
        try:
            with open(csv_path, 'r', encoding='gbk', newline='') as file:
                csv_reader = csv.DictReader(file)
                for row in csv_reader:
                    data_rows.append(row)
        except UnicodeDecodeError:
            with open(csv_path, 'r', encoding='utf-8', errors='ignore', newline='') as file:
                csv_reader = csv.DictReader(file)
                for row in csv_reader:
                    data_rows.append(row)
    
    # 转换为DataFrame以保持兼容性
    df = pd.DataFrame(data_rows)
    
    # 处理BOM标记问题
    if '\ufeffheading_1' in df.columns:
        df = df.rename(columns={'\ufeffheading_1': 'heading_1'})
    
    # 将空字符串转换为NaN以便统一处理
    df = df.replace('', pd.NA)
    
    print(f"文件列名: {list(df.columns)}")
    print(f"数据形状: {df.shape}")
    print(f"前5行数据预览:")
    print(df.head())
    print(f"非空行统计:")
    for col in ['heading_1', 'heading_2', 'title', 'content']:
        if col in df.columns:
            non_empty = df[col].notna().sum()
            print(f"  {col}: {non_empty} 行非空")
    
    # 查找包含内容的列
    content_column = None
    if 'content' in df.columns:
        content_column = 'content'
    else:
        # 查找可能包含URL或内容的列
        for col in df.columns:
            if any(keyword in col.lower() for keyword in ['url', 'content', 'link', '链接', '内容']):
                content_column = col
                break
        
        # 如果还没找到，使用最后一列（通常包含URL）
        if content_column is None:
            content_column = df.columns[-1]
    
    print(f"使用列: {content_column}")
    
    # 处理所有行，不仅仅是有content的行
    print(f"总行数: {len(df)}")
    
    # 获取有内容的行进行内容解析
    content_series = df[content_column].dropna()
    content_list = []
    for item in content_series:
        try:
            if pd.isna(item):
                content_list.append("")
            else:
                content_list.append(str(item))
        except (ValueError, TypeError):
            # 如果pd.isna()失败（比如对于数组），直接转换为字符串
            content_list.append(str(item))
    print(f"有content的行数: {len(content_list)}")
    
    # 处理有内容的行
    contents = parse_contents(content_list) if content_list else []
    content_index = 0
    
    # 创建结果记录 - 处理所有行
    records = []
    for i in range(len(df)):
        record = df.iloc[i].to_dict()
        
        # 处理NaN值，将其转换为空字符串或合适的默认值
        for key, value in record.items():
            try:
                if pd.isna(value):
                    record[key] = ""
                elif isinstance(value, float) and key != 'content':
                    # 如果是数字但不是content列，转换为字符串
                    try:
                        record[key] = str(value) if not pd.isna(value) else ""
                    except (ValueError, TypeError):
                        record[key] = str(value)
            except (ValueError, TypeError):
                # 如果pd.isna()失败（比如对于数组），直接处理
                if isinstance(value, float) and key != 'content':
                    record[key] = str(value)
                elif value is None:
                    record[key] = ""
                else:
                    record[key] = str(value)
        
        # 为有content的行分配处理后的内容
        try:
            has_content = pd.notna(df.iloc[i][content_column])
        except (ValueError, TypeError):
            # 如果pd.notna()失败，检查是否为None
            has_content = df.iloc[i][content_column] is not None
        
        if has_content and content_index < len(contents):
            record['content'] = contents[content_index]
            content_index += 1
        else:
            record['content'] = ""  # 没有content的行设置为空字符串
            
        records.append(record)
    
    return records

def csv_to_word(csv_path):
    doc_path = export_to_word(extract_from_csv(csv_path))
    append_company_info_and_disclaimer(doc_path)
    return doc_path


if __name__ == "__main__":
    csv_path = r"temp-data/国资委9月15刊排版 .csv"
    doc_path = csv_to_word(csv_path)
    print(doc_path)
