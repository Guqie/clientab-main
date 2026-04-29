# -*- coding: utf-8 -*-
"""
战新与未来产业月报第四期 - 精确修复脚本
"""
from docx import Document
import re

def read_and_fix():
    doc = Document(r'd:\桌面\战新与未来产业月报第四期0427.docx')
    
    fixes = []
    
    # 定义所有需要修复的内容
    repairs = [
        # 标题格式问题
        (r'^3\.(\S)', r'3. \1', '标题序号后加空格'),
        (r'^6\.(\S)', r'6. \1', '标题序号后加空格'),
        (r'^7\.(\S)', r'7. \1', '标题序号后加空格'),
        (r'^5\.(\S)', r'5. \1', '标题序号后加空格'),
        (r'^9\.(\S)', r'9. \1', '标题序号后加空格'),
    ]
    
    # 精确修复
    precise_fixes = {
        '工信部表示动态跟踪': '工信部将动态跟踪',
        '液氢被认为是实现': '液氢是实现',
        '介绍，因极高的能量密度': '因极高的能量密度',
        '水质分析实验室中': '水质检测实验室中',
        '偏差分析、高级星座设计': '偏差评估与高级星座设计',
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 跳过空段落
        if not text.strip():
            continue
        
        new_text = text
        modified = False
        
        # 1. 处理标题格式 (数字.后面不是空格的情况)
        if re.match(r'^\d+\.', text):
            # 检查数字序号后是否需要加空格
            match = re.match(r'^(\d+\.)(\S)', text)
            if match:
                new_text = match.group(1) + ' ' + match.group(2)
                # 检查是否还有其他内容
                rest = text[len(match.group(0)):]
                if rest:
                    new_text += rest
                if new_text != text:
                    fixes.append(f"[行{i+1}] 标题格式修复")
                    modified = True
        
        # 2. 处理精确替换
        for old, new in precise_fixes.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
                fixes.append(f"[行{i+1}] '{old}' → '{new}'")
                modified = True
        
        # 3. 更新段落
        if modified and new_text != text:
            # 保留原有格式属性
            if para.runs:
                first_run = para.runs[0]
                # 保存样式信息
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                
                # 清空并重新设置
                for run in para.runs:
                    run.text = ''
                para.runs[0].text = new_text
                
                # 恢复样式
                if font_name:
                    para.runs[0].font.name = font_name
                if font_size:
                    para.runs[0].font.size = font_size
                if bold is not None:
                    para.runs[0].font.bold = bold
    
    # 保存文件
    output_path = r'd:\桌面\战新与未来产业月报第四期0427_已修复.docx'
    doc.save(output_path)
    
    return fixes, output_path

def main():
    print("=" * 60)
    print("战新与未来产业月报第四期 - 精确修复")
    print("=" * 60)
    
    fixes, output_path = read_and_fix()
    
    print(f"\n共完成 {len(fixes)} 处修改:\n")
    
    for fix in fixes:
        print(f"  {fix}")
    
    print(f"\n文件已保存到:")
    print(f"  {output_path}")
    
    # 验证修复
    print("\n" + "=" * 60)
    print("验证修复结果...")
    print("=" * 60)
    
    doc = Document(output_path)
    
    # 验证标题格式
    title_issues = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^\d+\.', text):
            if text[1] == '.' and text[2] != ' ':
                title_issues += 1
                print(f"  [标题问题] 行{i+1}: {text[:50]}")
    
    if title_issues == 0:
        print("  ✓ 标题格式问题已全部修复")
    
    # 验证关键词
    check_words = ['表示', '认为', '介绍', '偏差分析', '水质分析实验室']
    word_issues = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for word in check_words:
            if word in text:
                word_issues += 1
                print(f"  [关键词] 行{i+1}: 仍含'{word}'")
    
    if word_issues == 0:
        print("  ✓ 关键词问题已全部修复")
    
    return fixes

if __name__ == '__main__':
    main()
