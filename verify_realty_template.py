#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
房地产模板验证脚本
验证图片格式设置（居中对齐、首行缩进0）和^l转^p换行符替换功能
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

def check_realty_document(doc_path):
    """
    检查房地产模板生成的Word文档
    
    Args:
        doc_path (str): 文档路径
        
    Returns:
        dict: 检查结果统计
    """
    print(f"正在检查文档: {doc_path}")
    
    if not os.path.exists(doc_path):
        print(f"错误：文档不存在 - {doc_path}")
        return None
    
    try:
        doc = Document(doc_path)
        print(f"文档加载成功，共 {len(doc.paragraphs)} 个段落")
    except Exception as e:
        print(f"错误：无法加载文档 - {e}")
        return None
    
    # 统计变量初始化
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'image_paragraphs': 0,
        'centered_images': 0,
        'zero_indent_images': 0,
        'l_characters': 0,
        'p_characters': 0,
        'problematic_images': []
    }
    
    print("\n=== 开始段落分析 ===")
    
    # 遍历所有段落
    for i, paragraph in enumerate(doc.paragraphs):
        # 检查段落是否包含图片
        has_image = False
        for run in paragraph.runs:
            # 检查run中是否包含图片元素
            if run._element.xpath('.//a:blip'):
                has_image = True
                break
            # 备用检测方法：检查drawing元素
            if run._element.xpath('.//w:drawing'):
                has_image = True
                break
        
        if has_image:
            stats['image_paragraphs'] += 1
            
            # 检查对齐方式
            alignment = paragraph.alignment
            if alignment == WD_ALIGN_PARAGRAPH.CENTER or alignment == 1:
                stats['centered_images'] += 1
            else:
                stats['problematic_images'].append({
                    'paragraph_index': i,
                    'alignment': alignment,
                    'alignment_name': get_alignment_name(alignment),
                    'text': paragraph.text[:50] if paragraph.text else "(空段落)"
                })
            
            # 检查首行缩进
            first_line_indent = paragraph.paragraph_format.first_line_indent
            if first_line_indent is None or first_line_indent.inches == 0:
                stats['zero_indent_images'] += 1
            
            print(f"图片段落 #{stats['image_paragraphs']}: 索引{i}, 对齐={get_alignment_name(alignment)}, 首行缩进={first_line_indent}")
        
        # 统计^l和^p字符
        text = paragraph.text
        stats['l_characters'] += text.count('^l')
        stats['p_characters'] += text.count('^p')
    
    return stats

def get_alignment_name(alignment):
    """
    获取对齐方式的名称
    
    Args:
        alignment: Word对齐方式枚举值
        
    Returns:
        str: 对齐方式名称
    """
    alignment_map = {
        None: "None",
        0: "LEFT",
        1: "CENTER", 
        2: "RIGHT",
        3: "JUSTIFY",
        4: "DISTRIBUTE"
    }
    return alignment_map.get(alignment, f"UNKNOWN({alignment})")

def print_verification_results(stats):
    """
    打印验证结果
    
    Args:
        stats (dict): 统计结果
    """
    if not stats:
        return
    
    print("\n" + "="*60)
    print("房地产模板验证结果")
    print("="*60)
    
    print(f"总段落数: {stats['total_paragraphs']}")
    print(f"图片段落数: {stats['image_paragraphs']}")
    print(f"居中对齐的图片: {stats['centered_images']}")
    print(f"首行缩进为0的图片: {stats['zero_indent_images']}")
    print(f"^l字符数: {stats['l_characters']}")
    print(f"^p字符数: {stats['p_characters']}")
    
    print("\n=== 功能验证 ===")
    
    # 图片居中对齐验证
    if stats['image_paragraphs'] > 0:
        center_rate = (stats['centered_images'] / stats['image_paragraphs']) * 100
        print(f"✓ 图片居中对齐: {stats['centered_images']}/{stats['image_paragraphs']} ({center_rate:.1f}%)")
        if center_rate < 100:
            print("  ⚠️  存在未居中对齐的图片:")
            for img in stats['problematic_images']:
                print(f"    - 段落{img['paragraph_index']}: 对齐方式={img['alignment_name']}, 文本='{img['text']}'")
    else:
        print("ℹ️  未发现图片段落")
    
    # 首行缩进验证
    if stats['image_paragraphs'] > 0:
        indent_rate = (stats['zero_indent_images'] / stats['image_paragraphs']) * 100
        print(f"✓ 图片首行缩进为0: {stats['zero_indent_images']}/{stats['image_paragraphs']} ({indent_rate:.1f}%)")
    
    # 换行符转换验证
    if stats['l_characters'] == 0 and stats['p_characters'] == 0:
        print("✓ ^l转^p功能: 正确 (未发现^l或^p字符)")
    elif stats['l_characters'] == 0:
        print(f"✓ ^l转^p功能: 正确 (发现{stats['p_characters']}个^p字符，无^l字符)")
    else:
        print(f"❌ ^l转^p功能: 异常 (仍有{stats['l_characters']}个^l字符未转换)")
    
    print("\n=== 测试总结 ===")
    
    issues = []
    if stats['image_paragraphs'] > 0 and stats['centered_images'] < stats['image_paragraphs']:
        issues.append(f"图片居中对齐不完整 ({stats['centered_images']}/{stats['image_paragraphs']})")
    
    if stats['image_paragraphs'] > 0 and stats['zero_indent_images'] < stats['image_paragraphs']:
        issues.append(f"图片首行缩进设置不完整 ({stats['zero_indent_images']}/{stats['image_paragraphs']})")
    
    if stats['l_characters'] > 0:
        issues.append(f"^l转^p转换不完整 (剩余{stats['l_characters']}个^l字符)")
    
    if not issues:
        print("🎉 所有功能验证通过！")
    else:
        print("⚠️  发现以下问题:")
        for issue in issues:
            print(f"   - {issue}")

def main():
    """主函数"""
    # 房地产模板生成的文档路径
    doc_path = "temp-data/1758620917_204430_realty.docx"
    
    print("房地产模板功能验证")
    print("="*60)
    print("验证项目:")
    print("1. 图片格式设置 - 居中对齐")
    print("2. 图片格式设置 - 首行缩进0")
    print("3. ^l转^p换行符替换功能")
    print("="*60)
    
    # 执行验证
    stats = check_realty_document(doc_path)
    
    # 打印结果
    print_verification_results(stats)

if __name__ == "__main__":
    main()