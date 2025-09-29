from docx import Document
import re

def detailed_document_check():
    """详细检查Word文档中的所有内容"""
    doc_path = 'temp-data/1757923949_523631_guoziwei.docx'
    
    try:
        doc = Document(doc_path)
        print(f"文档总段落数: {len(doc.paragraphs)}")
        
        # 查找所有包含日期格式的段落
        date_patterns = [
            r'2025/\d{1,2}/\d{1,2}',  # 2025/9/11格式
            r'2025-\d{1,2}-\d{1,2}',  # 2025-09-11格式
            r'\d{4}/\d{1,2}/\d{1,2}', # 任意年份/月/日格式
        ]
        
        found_dates = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                for pattern in date_patterns:
                    matches = re.findall(pattern, text)
                    if matches:
                        found_dates.append((i, text, matches))
        
        print(f"\n找到 {len(found_dates)} 个包含日期的段落:")
        for i, text, matches in found_dates:
            print(f"段落{i}: {text}")
            print(f"  匹配的日期: {matches}")
        
        # 查找可能的Times New Roman字体段落
        times_new_roman_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # 检查段落中的runs
                for run in paragraph.runs:
                    if run.font.name == "Times New Roman":
                        times_new_roman_paragraphs.append((i, text, run.text))
                        break
        
        print(f"\n找到 {len(times_new_roman_paragraphs)} 个Times New Roman字体段落:")
        for i, text, run_text in times_new_roman_paragraphs:
            print(f"段落{i}: {text}")
            print(f"  Times New Roman文本: {run_text}")
        
        # 搜索特定的日期值
        target_dates = ['2025/9/11', '2025/9/10', '2025/9/8', '2025/9/2', '2025/9/4', '2025/9/5', '2025/9/15', '2025/9/9', '2025/9/6']
        
        for target_date in target_dates:
            found_in_paragraphs = []
            for i, paragraph in enumerate(doc.paragraphs):
                if target_date in paragraph.text:
                    found_in_paragraphs.append(i)
            
            if found_in_paragraphs:
                print(f"\n日期 {target_date} 出现在段落: {found_in_paragraphs}")
                for para_idx in found_in_paragraphs:
                    print(f"  段落{para_idx}: {doc.paragraphs[para_idx].text.strip()}")
        
        # 检查文档中间部分的段落（可能日期被添加在内容中间）
        print("\n检查文档中间部分的段落（第100-120段）:")
        for i in range(100, min(120, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if text and len(text) < 50:  # 短段落可能是日期
                print(f"段落{i}: {text}")
    
    except Exception as e:
        print(f"检查文档时出错: {e}")

if __name__ == "__main__":
    detailed_document_check()