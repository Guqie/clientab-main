# -*- coding: utf-8 -*-
"""
验证已修复的文档
"""
from docx import Document
import re

def verify_fixed():
    doc = Document(r'd:\桌面\战新与未来产业月报第四期0427_已修复.docx')
    
    print("=" * 60)
    print("已修复文档 - 最终验证报告")
    print("=" * 60)
    
    issues = []
    
    # 1. 检查标题格式
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^\d+\.', text):
            if text[1] == '.' and len(text) > 2 and text[2] != ' ':
                issues.append(f"【标题格式】行{i+1}: {text[:60]}")
    
    # 2. 检查必须删除的词
    must_delete = [
        ('表示', '工信部表示'),
        ('介绍，', '据悉，'),
        ('被認為', '液氢被認為'),
        ('水质分析实验室', '水质分析实验室'),
        ('偏差分析', '偏差分析'),
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # 工信部表示
        if '工信部表示' in text:
            issues.append(f"【用词】行{i+1}: 含'工信部表示'")
        # 液氢被认为
        if '液氢被認為' in text or '液氢被认为' in text:
            issues.append(f"【用词】行{i+1}: 含'液氢被认为'")
        # 介绍，
        if '介绍，' in text:
            issues.append(f"【用词】行{i+1}: 含'介绍，'")
        # 水质分析实验室
        if '水质分析实验室' in text:
            issues.append(f"【用词】行{i+1}: 含'水质分析实验室'")
        # 偏差分析
        if '偏差分析' in text and '偏差评估' not in text:
            issues.append(f"【用词】行{i+1}: 含'偏差分析'")
    
    # 3. 检查中车
    zhongche_words = ['中国中车', 'CRRC', '中车集团', '株机', '株洲所']
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for word in zhongche_words:
            if word in text:
                issues.append(f"【中车】行{i+1}: 含'{word}'")
    
    print(f"\n文档总段落数: {len(doc.paragraphs)}")
    
    if issues:
        print(f"\n发现 {len(issues)} 处问题:\n")
        for issue in issues:
            print(f"  {issue}")
        print("\n[FAIL] 文档仍需修改")
    else:
        print("\n[PASS] 所有检查项通过！")
        print("\n文档已准备就绪，可以交付。")
    
    return len(issues) == 0

if __name__ == '__main__':
    verify_fixed()
