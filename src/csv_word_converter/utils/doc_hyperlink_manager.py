#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import logging
import random
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional

# 导入WPS兼容性优化的函数
try:
    from ..utils.doc_utils import add_internal_hyperlink
except ImportError:
    # 回退到utils目录的导入
    import sys
    from pathlib import Path
    sys.path.append(str(Path(__file__).parent.parent.parent.parent / "utils"))
    from doc_utils import add_internal_hyperlink

class DocumentHyperlinkManager:
    """
    文档超链接管理器，负责Word文档中书签和内部超链接的创建与管理。
    """

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        if not self.logger.handlers:
            logging.basicConfig(level=logging.INFO)

    def _add_bookmark(self, paragraph, bookmark_name: str, visible: bool = False):
        """
        在段落开头添加书签（WPS兼容性优化版本）。

        Args:
            paragraph: docx.paragraph.Paragraph 对象。
            bookmark_name: 书签名称。
            visible: 是否在书签位置显示可见标记（默认为False，隐藏书签）。
        """
        self.logger.info(f"在段落开头添加隐藏书签: {bookmark_name}")
        
        # 生成唯一的书签ID（使用时间戳确保唯一性）
        import time
        bookmark_id = str(int(time.time() * 1000) % 1000000)
        
        # 创建书签起始标签
        start_tag = OxmlElement('w:bookmarkStart')
        start_tag.set(qn('w:id'), bookmark_id)
        start_tag.set(qn('w:name'), bookmark_name)
        
        # 创建书签结束标签（注意：结束标签不应设置name属性，这是WPS兼容性关键）
        end_tag = OxmlElement('w:bookmarkEnd')
        end_tag.set(qn('w:id'), bookmark_id)  # 只设置ID，不设置name
        
        # 如果需要可见，则在书签位置添加一个可见的Run
        if visible:
            # 创建一个Run来容纳可见的标记文本
            run_for_text = OxmlElement('w:r')
            # 创建文本元素
            text_element = OxmlElement('w:t')
            text_element.text = f'[{bookmark_name.upper()}]' # 使用大写形式以示区别
            run_for_text.append(text_element)
            # 将Run插入到段落的开头
            paragraph._p.insert(0, run_for_text)
            # 调整书签标签位置
            paragraph._p.insert(0, start_tag)
            paragraph._p.insert(2, end_tag)
        else:
            # 隐藏书签：直接在段落开头插入书签标签
            paragraph._p.insert(0, start_tag)
            paragraph._p.insert(1, end_tag)

    def _replace_with_hyperlink(self, paragraph, hyperlink_text: str, anchor_name: str):
        """
        将段落的全部内容替换为一个内部超链接（使用WPS兼容性优化版本）。

        Args:
            paragraph: docx.paragraph.Paragraph 对象。
            hyperlink_text: 超链接显示的文本。
            anchor_name: 目标书签的名称。
        """
        # 清空段落现有内容
        paragraph.clear()
        
        # 使用WPS兼容性优化的add_internal_hyperlink函数
        add_internal_hyperlink(
            paragraph=paragraph,
            link_to=anchor_name,
            text=hyperlink_text,
            tooltip="点击返回目录",
            font_name="宋体",
            font_size=12,
            underline=True,
            font_color="#0563C1"  # 标准超链接蓝色
        )

    def create_return_to_toc_hyperlinks(self, doc_path: str, target_keyword: str, placeholder_text: str):
        """
        遍历文档，将所有指定占位符替换为指向动态书签的超链接。
        该书签的位置由配置文件中的 'target_bookmark' 关键词决定。

        Args:
            doc_path: Word文档的路径。
            target_keyword: 用于定位目录书签的关键词。
            placeholder_text: 需要被替换为超链接的占位符文本。
        """
        self.logger.info("开始创建“返回目录”的内部超链接...")
        try:
            doc = Document(doc_path)
            if not doc.paragraphs:
                self.logger.warning("文档为空，无法创建超链接。")
                return

            toc_anchor_name = "TOC_ANCHOR"
            target_paragraph = None
            
            # 遍历段落查找关键词
            for para in doc.paragraphs:
                if target_keyword in para.text:
                    target_paragraph = para
                    break
            
            # 如果在段落中没找到，则遍历表格
            if not target_paragraph:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if target_keyword in para.text:
                                    target_paragraph = para
                                    break
                            if target_paragraph: break
                        if target_paragraph: break
                    if target_paragraph: break

            # 根据是否找到关键词来设置书签
            if target_paragraph:
                self._add_bookmark(target_paragraph, toc_anchor_name, visible=False)  # 修改为隐藏书签
                self.logger.info(f"在关键词 '{target_keyword}' 所在位置创建了隐藏书签: '{toc_anchor_name}'")
            else:
                # 回退机制：在文档顶部创建书签
                first_paragraph = doc.paragraphs[0]
                self._add_bookmark(first_paragraph, toc_anchor_name, visible=False)  # 修改为隐藏书签
                self.logger.warning(f"未在文档中找到关键词 '{target_keyword}'，已在文档开头创建隐藏书签: '{toc_anchor_name}'")

            # 遍历并替换所有占位符为超链接
            count = 0
            # 遍历正文段落
            for para in doc.paragraphs:
                if para.text.strip() == placeholder_text:
                    self._replace_with_hyperlink(para, placeholder_text, toc_anchor_name)
                    count += 1

            # 遍历表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip() == placeholder_text:
                                self._replace_with_hyperlink(para, placeholder_text, toc_anchor_name)
                                count += 1
            
            doc.save(doc_path)
            self.logger.info(f"成功将 {count} 个“{placeholder_text}”占位符转换为超链接。")

        except Exception as e:
            self.logger.error(f"创建超链接时发生错误: {e}", exc_info=True)