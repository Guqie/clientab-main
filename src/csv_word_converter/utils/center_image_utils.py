#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
图片描述段落居中处理工具

该模块提供图片描述段落的居中处理功能，用于文档后处理阶段。
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH


def center_image_description_paragraphs(doc):
    """
    将文档中的图片段落和图片描述段落居中对齐，并清理格式问题
    
    该函数遍历文档中的所有段落，识别包含图片的段落（纯图片段落）和包含图片描述的段落，
    并将其设置为居中对齐。同时清理可能存在的格式问题，如^l换行符等。
    
    Args:
        doc: python-docx Document对象，要处理的Word文档
        
    Returns:
        None: 直接修改传入的文档对象
        
    Note:
        - 该函数会识别包含实际图片的段落（通过检查run中的图片元素）
        - 同时识别包含"图"、"Figure"、"图片"等关键词的图片描述段落
        - 仅对符合条件的段落进行居中处理，不影响其他段落格式
        - 处理过程中会保留段落的其他格式属性
        - 清理图片段落中的^l换行符和多余的空白run
    """
    if not doc:
        return
        
    try:
        # 定义图片描述的关键词
        image_keywords = ["图", "Figure", "图片", "Fig.", "图表", "插图"]
        
        # 遍历文档中的所有段落
        for paragraph in doc.paragraphs:
            # 检查段落是否包含实际图片（通过检查run中的图片元素）
            has_image = any(run._element.xpath('.//pic:pic') for run in paragraph.runs)
            
            # 检查段落是否包含图片描述关键词
            text = paragraph.text.strip()
            is_image_description = text and any(keyword in text for keyword in image_keywords)
            
            # 如果是图片段落或图片描述段落，设置为居中对齐
            if has_image or is_image_description:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 确保图片段落的缩进也正确设置
                if has_image:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.first_line_indent = 0
                    paragraph_format.left_indent = 0
                    paragraph_format.right_indent = 0
                    
                    # 清理图片段落中的问题：删除空白run和包含^l的run
                    runs_to_remove = []
                    for run in paragraph.runs:
                        # 检查run是否只包含空白字符或^l换行符
                        run_text = run.text
                        if run_text and (run_text.strip() == '' or '^l' in run_text or '\f' in run_text):
                            # 如果run不包含图片，则标记为删除
                            if not run._element.xpath('.//pic:pic'):
                                runs_to_remove.append(run)
                    
                    # 删除标记的run
                    for run in runs_to_remove:
                        try:
                            run._element.getparent().remove(run._element)
                        except Exception as e:
                            print(f"删除问题run时出错: {e}")
                
        # 同时处理表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue
                            
                        text = paragraph.text.strip()
                        is_image_description = any(keyword in text for keyword in image_keywords)
                        
                        if is_image_description:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
    except Exception as e:
        print(f"处理图片描述段落居中时出错: {e}")