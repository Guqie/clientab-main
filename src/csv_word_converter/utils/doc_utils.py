"""
Doc Utilities (适配层)

作用：
- 统一对外暴露 Word/文本处理相关的通用函数；
- 对 csv_to_word 的直接依赖收敛在 utils 层，降低上层业务（如 universal_csv_to_word.py）的跨模块耦合；
- 若未来替换实现（例如迁移到 document_generator.py 或新的文本规整库），仅需在此文件适配即可，调用方无需改动。
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# 当前阶段采用"转发导出"，保持与历史实现完全兼容
# 后续如需替换底层实现，可在此处改为自有实现或桥接到其他模块

# 导入图片处理工具
from .center_image_utils import center_image_description_paragraphs


__all__ = [
    "process_lines",

    "add_hyperlinks_post_processing",
    "replace_halfwidth_quotes_with_fullwidth",
    "remove_special_symbols",
    "change_digits_letters_punctuation_to_times_new_roman",
    "remove_space_between_chinese_and_digits_letters_punctuation",
    "center_image_description_paragraphs",
    "convert_to_fullwidth",
    "parse_content",
    "normalize_spaces_and_convert_punct_except_period",
    "add_internal_hyperlink",

    # 新增通用工具导出
    "apply_paragraph_format",
    "compute_heading_level",
    "format_title_text",
    "add_return_directory_placeholder",
    "add_bookmark_to_paragraph_xml",
    "create_target_bookmark_by_keyword",
    "convert_return_placeholders_to_hyperlinks",
]


def apply_paragraph_format(paragraph, config: dict):
    """应用段落格式到给定段落。

    该工具函数将通用的段落格式配置映射到 python-docx 的段落对象上，避免重复实现。
    支持的配置键：
    - alignment: 对齐方式（left/center/right/justify）
    - first_line_indent: 首行缩进（以“字符数”表示，内部换算为 Pt(字符数*12)）
    - line_spacing: 行距（数值，1.0/1.5/2 等）
    - space_after: 段后间距（以磅为单位）

    Args:
        paragraph: python-docx 段落对象
        config: 字典形式的段落格式配置
    """
    if not config:
        return
    # 对齐方式
    alignment = config.get("alignment")
    if alignment:
        alignment_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    # 首行缩进（字符 -> 磅）
    if "first_line_indent" in config:
        indent_chars = config["first_line_indent"]
        try:
            paragraph.paragraph_format.first_line_indent = Pt(int(indent_chars) * 12)
        except Exception:
            paragraph.paragraph_format.first_line_indent = Pt(24)
    # 行距
    if "line_spacing" in config:
        paragraph.paragraph_format.line_spacing = config["line_spacing"]
    # 段后间距
    if "space_after" in config:
        paragraph.paragraph_format.space_after = Pt(config["space_after"])


def compute_heading_level(heading_type: str, title_mapping: dict | None = None) -> int | None:
    """根据标题类型与映射表，计算 Word 标题级别。

    Args:
        heading_type: 标题类型，如 "heading_1"/"heading_2"/"heading_3"/"title"
        title_mapping: 可选的映射表，如 {"heading_1":1, "heading_2":2}
    Returns:
        int | None: 对应的级别；若无法识别则返回 None（调用方可回退为普通段落）。
    """
    mapping = title_mapping or {}
    if heading_type in mapping:
        return mapping[heading_type]
    default_map = {
        "heading_1": 1,
        "heading_2": 2,
        "heading_3": 3,
        "title": 3,
    }
    return default_map.get(heading_type)


def format_title_text(text: str, heading_type: str) -> str:
    """根据标题类型格式化文本：对 title 类型增加【】包裹。"""
    return f"【{text}】" if heading_type == "title" else text


def add_return_directory_placeholder(doc, return_link_config: dict | None = None):
    """添加“返回目录”占位符段落。

    说明：该函数仅负责创建带有指定文本与段落/字体格式的占位段落，不负责添加超链接；
    超链接应在文档生成完成后的后处理阶段统一添加。

    Args:
        doc: python-docx Document 对象
        return_link_config: 配置项，支持 text/font_name/font_size/alignment/underline
    """
    cfg = return_link_config or {}
    from docx.enum.text import WD_UNDERLINE

    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    alignment = cfg.get("alignment", "right")
    if alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(cfg.get("text", "返回目录"))
    run.font.name = cfg.get("font_name", "宋体")
    run.font.size = Pt(cfg.get("font_size", 12))
    if cfg.get("underline", False):
        run.font.underline = WD_UNDERLINE.SINGLE


def add_bookmark_to_paragraph_xml(paragraph, bookmark_name: str):
    """在段落中插入书签的起止标记（bookmarkStart/bookmarkEnd），不改变段落文本。

    与 csv_to_word.add_bookmark 的区别：本函数不会新增文本 run，仅在段落 XML 结构中插入书签标记，
    因此对原有排版影响最小，适用于“对现有标题/目录文本打书签”的场景。

    Args:
        paragraph: python-docx 段落对象
        bookmark_name: 书签名称
    """
    # 生成稳定的书签 ID（避免与其他书签冲突）
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
    # 构造 bookmarkStart 与 bookmarkEnd 元素
    start = paragraph._element.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart"
    )
    start.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", bookmark_id)
    start.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name", bookmark_name)
    end = paragraph._element.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd")
    end.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", bookmark_id)
    # 插入位置：在第一个 run 之前插入 start，段落末尾追加 end
    insert_index = 0
    for i, child in enumerate(paragraph._element):
        if child.tag.endswith("r"):
            insert_index = i
            break
        elif child.tag.endswith("pPr"):
            insert_index = i + 1
    paragraph._element.insert(insert_index, start)
    paragraph._element.append(end)


def create_target_bookmark_by_keyword(doc, target_text: str) -> bool:
    """查找包含指定关键字的段落，并为该段落插入书签（不改变原样式）。

    Args:
        doc: Document 对象
        target_text: 需要定位并打书签的关键字（如“目录”）
    Returns:
        bool: 若成功创建至少一个书签则为 True，否则为 False。
    """
    try:
        # 先检查是否已有同名书签，避免重复创建引发 WPS 跳转失败
        existing = _get_existing_bookmark_names(doc)
        clean_name = str(target_text).strip().replace(" ", "_")
        if target_text in existing or clean_name in existing:
            return True
        for paragraph in doc.paragraphs:
            if target_text in (paragraph.text or ""):
                add_bookmark_to_paragraph_xml(paragraph, target_text)
                return True
        # 也可扩展到表格内查找，这里保持最小实现以通用性优先
        return False
    except Exception:
        return False


def convert_return_placeholders_to_hyperlinks(doc, placeholder_text="返回目录", bookmark_name="toc_bookmark"):
    """
    将文档中所有占位符文本转换为内部超链接，指向指定的书签。

    Args:
        doc (docx.Document): 要处理的文档对象。
        placeholder_text (str): 要查找和替换的占位符文本。
        bookmark_name (str): 目标书签的名称。
    """
    for p in doc.paragraphs:
        # 检查段落中是否包含占位符文本
        if placeholder_text in p.text:
            # 清空段落，准备重建
            p.clear()
            # 直接在段落上添加内部超链接，传入正确的书签和文本
            add_internal_hyperlink(p, bookmark_name, placeholder_text)


def add_internal_hyperlink(
    paragraph, link_to, text, tooltip=None, font_name="宋体", font_size=12, underline=True, font_color=None
):
    """在段落中添加指向文档内部书签的超链接（WPS兼容性增强版本）

    Args:
        paragraph: 段落对象
        link_to: 目标书签名称
        text: 显示文本
        tooltip: 鼠标悬停提示
        font_name: 字体名称
        font_size: 字体大小（磅）
        underline: 是否下划线
        font_color: 字体颜色
    """
    from docx.oxml.shared import OxmlElement, qn

    # 创建超链接元素
    hyperlink = OxmlElement("w:hyperlink")

    # 设置锚点属性（指向书签）- 确保书签名称格式正确
    clean_anchor = str(link_to).strip()
    hyperlink.set(qn("w:anchor"), clean_anchor)

    # 设置工具提示
    if tooltip:
        hyperlink.set(qn("w:tooltip"), str(tooltip))

    # WPS兼容性增强：添加历史属性（WPS可能需要此属性）
    hyperlink.set(qn("w:history"), "1")

    # 创建运行元素
    new_run = OxmlElement("w:r")

    # 创建运行属性 - 增强WPS兼容性
    rPr = OxmlElement("w:rPr")

    # 设置字体 - 明确指定所有字体属性
    font_element = OxmlElement("w:rFonts")
    font_element.set(qn("w:ascii"), font_name)
    font_element.set(qn("w:eastAsia"), font_name)
    font_element.set(qn("w:hAnsi"), font_name)
    font_element.set(qn("w:cs"), font_name)  # 添加复杂脚本字体
    rPr.append(font_element)

    # 设置字体大小（Word中w:sz为磅值的两倍）
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)

    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(szCs)

    # 设置下划线
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    # 设置字体颜色 - 优化颜色处理
    color = OxmlElement("w:color")
    if font_color:
        # 处理各种颜色格式
        hex_val = _parse_color_to_hex(font_color)
        if hex_val:
            color.set(qn("w:val"), hex_val)
        else:
            color.set(qn("w:themeColor"), "hyperlink")
    else:
        # 使用主题超链接颜色，增强兼容性
        color.set(qn("w:themeColor"), "hyperlink")
    rPr.append(color)

    # 添加超链接样式 - WPS兼容性关键
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # WPS兼容性增强：添加语言属性
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    rPr.append(lang)

    # 将运行属性添加到运行
    new_run.append(rPr)

    # 创建文本元素
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = str(text)
    new_run.append(text_element)

    # 将运行添加到超链接
    hyperlink.append(new_run)

    # 将超链接添加到段落
    paragraph._element.append(hyperlink)


def _parse_color_to_hex(font_color):
    """解析各种颜色格式为十六进制值"""
    import re

    if not font_color:
        return None

    color_map = {"red": "FF0000", "blue": "0000FF", "black": "000000", "green": "008000", "purple": "800080", "darkred": "8B0000"}  # 添加深红色

    # 颜色名称
    if isinstance(font_color, str) and font_color.lower() in color_map:
        return color_map[font_color.lower()]

    # #RRGGBB 或 RRGGBB
    if isinstance(font_color, str) and re.fullmatch(r"#?[0-9A-Fa-f]{6}", font_color.strip()):
        return font_color.strip().lstrip("#").upper()

    # rgb(r,g,b)
    if isinstance(font_color, str) and font_color.strip().lower().startswith("rgb"):
        nums = re.findall(r"\d+", font_color)
        if len(nums) == 3:
            r_val, g_val, b_val = [min(255, max(0, int(n))) for n in nums]
            return f"{r_val:02X}{g_val:02X}{b_val:02X}"

    # (r,g,b) 或 [r,g,b]
    if isinstance(font_color, (tuple, list)) and len(font_color) == 3:
        r_val, g_val, b_val = [min(255, max(0, int(n))) for n in font_color]
        return f"{r_val:02X}{g_val:02X}{b_val:02X}"

    return None


# 新增：扫描文档已有书签名称，避免重复创建
def _get_existing_bookmark_names(doc) -> set:
    """返回文档中已存在的书签名称集合。

    目的：
    - 在创建“目录”等目标书签前，先复用已有书签，避免出现同名多书签导致 WPS 内部跳转异常。

    参数：
    - doc: python-docx Document 对象

    返回：
    - set[str]: 书签名称的集合
    """
    names = set()
    try:
        for elem in doc._element.iter():
            if isinstance(elem.tag, str) and elem.tag.endswith("bookmarkStart"):
                name_attr = elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
                if name_attr:
                    names.add(name_attr)
    except Exception:
        # 出现解析异常时返回已收集结果
        pass
    return names


def add_bookmark_to_paragraph_xml_enhanced(paragraph, bookmark_name: str):
    """增强版书签添加函数，提高WPS兼容性

    Args:
        paragraph: python-docx段落对象
        bookmark_name: 书签名称
    """
    import random
    import time

    # 生成更稳定的书签ID - 避免冲突
    timestamp = int(time.time() * 1000) % 1000000
    random_part = random.randint(1000, 9999)
    bookmark_id = str(timestamp + random_part)

    # 清理书签名称 - 确保符合规范
    clean_name = str(bookmark_name).strip().replace(" ", "_")

    # 构造bookmarkStart元素
    start = paragraph._element.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart"
    )
    start.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", bookmark_id)
    start.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name", clean_name)

    # 构造bookmarkEnd元素
    end = paragraph._element.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd")
    end.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", bookmark_id)

    # 插入位置优化 - 确保正确的XML结构
    insert_index = 0
    for i, child in enumerate(paragraph._element):
        if child.tag.endswith("r"):  # 找到第一个run元素
            insert_index = i
            break
        elif child.tag.endswith("pPr"):  # 段落属性之后
            insert_index = i + 1

    # 插入书签标记
    paragraph._element.insert(insert_index, start)
    paragraph._element.append(end)

    return bookmark_id


def create_target_bookmark_by_keyword_enhanced(doc, target_text: str) -> tuple[bool, str]:
    """增强版目标书签创建函数

    Args:
        doc: Document对象
        target_text: 目标关键字

    Returns:
        tuple[bool, str]: (是否成功, 实际书签名称)
    """
    try:
        # 先复用已有书签，避免重复创建
        existing = _get_existing_bookmark_names(doc)
        clean_name = str(target_text).strip().replace(" ", "_")
        if target_text in existing:
            return True, target_text
        if clean_name in existing:
            return True, clean_name

        # 首先尝试精确匹配
        for paragraph in doc.paragraphs:
            para_text = (paragraph.text or "").strip()
            if para_text == target_text:
                add_bookmark_to_paragraph_xml_enhanced(paragraph, target_text)
                return True, clean_name

        # 然后尝试包含匹配
        for paragraph in doc.paragraphs:
            para_text = (paragraph.text or "").strip()
            if target_text in para_text and para_text:
                add_bookmark_to_paragraph_xml_enhanced(paragraph, target_text)
                return True, clean_name

        # 最后在表格中查找
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_text = (paragraph.text or "").strip()
                        if target_text in para_text and para_text:
                            add_bookmark_to_paragraph_xml_enhanced(paragraph, target_text)
                            return True, clean_name

        return False, clean_name

    except Exception as e:
        print(f"创建书签时出错: {e}")
        return False, str(target_text).strip().replace(" ", "_")
