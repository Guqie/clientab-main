#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pandas as pd
import yaml
import logging
import re
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_UNDERLINE
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
from copy import deepcopy
import os
import time
import random
from pathlib import Path
from csv_word_converter.utils import (
    compute_heading_level,
    apply_paragraph_format,
    add_return_directory_placeholder,
    create_target_bookmark_by_keyword,
    convert_return_placeholders_to_hyperlinks,
    add_bookmark_to_paragraph_xml,
    compute_heading_level,
    format_title_text,
    add_internal_hyperlink,
    center_image_description_paragraphs,
)
from csv_word_converter.utils.doc_hyperlink_manager import DocumentHyperlinkManager


def apply_csv_field_mapping(df):
    """
    对CSV数据应用中文字段名到英文字段名的映射
    
    Args:
        df (pd.DataFrame): 原始CSV数据框
        
    Returns:
        pd.DataFrame: 映射后的数据框
    """
    field_mapping = {
        # 标题相关字段
        '标题': 'title',
        '一级栏目': 'heading_1', 
        '二级栏目': 'heading_2',
        '三级栏目': 'heading_3',
        
        # 内容相关字段
        '内容': 'content',
        '正文': 'content',
        '文章内容': 'content',
        
        # 来源和日期字段
        '来源': 'source',
        '信源': 'source',
        '日期': 'date',
        '发布日期': 'date',
        '时间': 'date',
        
        # URL字段
        'URL': 'url',
        '链接': 'url',
        '网址': 'url',
        
        # 其他可能的字段
        '序号': 'index',
        '编号': 'index',
        '分类': 'category',
        '类别': 'category',
    }
    
    # 创建列名映射字典（只映射存在的列）
    column_rename_dict = {}
    for chinese_name, english_name in field_mapping.items():
        if chinese_name in df.columns:
            column_rename_dict[chinese_name] = english_name
    
    # 应用列名映射
    if column_rename_dict:
        df = df.rename(columns=column_rename_dict)
        print(f"应用字段映射: {column_rename_dict}")
    
    return df

class DocumentTemplate(ABC):
    """文档模板抽象基类"""
    
    @abstractmethod
    def get_template_path(self) -> str:
        """获取模板文件路径"""
        pass
    
    @abstractmethod
    def get_style_config(self) -> dict:
        """获取样式配置"""
        pass
    
    @abstractmethod
    def process_content(self, item: dict) -> dict:
        """处理内容数据"""
        pass
    
    @abstractmethod
    def get_end_template_path(self) -> str:
        """获取结尾模板路径"""
        pass

class ConfigBasedTemplate(DocumentTemplate):
    """基于配置的模板实现"""
    
    def __init__(self, template_config: dict):
        self.config = template_config
    
    def get_template_path(self) -> str:
        return self.config['start_template']
    
    def get_style_config(self) -> dict:
        """获取样式配置，包括styles、target_bookmark、return_link、title_mapping等"""
        # 返回完整的配置，包括styles和其他配置项如target_bookmark、return_link、title_mapping等
        # 安全地处理styles配置，如果为None则使用空字典
        styles = self.config.get('styles', {})
        result = styles.copy() if styles is not None else {}
        
        # 添加模板级别的配置项
        if 'target_bookmark' in self.config:
            result['target_bookmark'] = self.config['target_bookmark']
        if 'return_link' in self.config:
            result['return_link'] = self.config['return_link']
        if 'title_mapping' in self.config:
            result['title_mapping'] = self.config['title_mapping']
        return result
    
    def process_content(self, item: dict) -> dict:
        """根据配置处理内容"""
        processed_item = item.copy()
        
        # 处理title格式化
        if 'title' in processed_item and processed_item['title']:
            title_config = self.config['styles'].get('title', {})
            if 'format' in title_config:
                processed_item['title'] = title_config['format'].format(processed_item['title'])
        
        return processed_item
    
    def get_end_template_path(self) -> str:
        return self.config.get('end_template', '')

class TemplateFactory:
    """模板工厂类"""
    
    def __init__(self, config_path: str = "templates_config.yaml"):
        self.config = self._load_config(config_path)
        self._validate_config()
    
    def _load_config(self, config_path: str) -> dict:
        """加载配置文件"""
        try:
            with open(config_path, 'r', encoding='utf-8') as file:
                return yaml.safe_load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"配置文件未找到: {config_path}")
        except yaml.YAMLError as e:
            raise ValueError(f"配置文件格式错误: {e}")
    
    def _validate_config(self):
        """验证配置文件"""
        if 'templates' not in self.config:
            raise ValueError("配置文件缺少 'templates' 字段")
        
        for template_name, template_config in self.config['templates'].items():
            required_keys = ['name', 'start_template', 'styles']
            for key in required_keys:
                if key not in template_config:
                    raise ValueError(f"模板 {template_name} 缺少必需字段: {key}")
    
    def create_template(self, template_type: str) -> DocumentTemplate:
        """创建模板实例"""
        if template_type not in self.config['templates']:
            available = list(self.config['templates'].keys())
            raise ValueError(f"不支持的模板类型: {template_type}，可用模板: {available}")
        
        template_config = self.config['templates'][template_type]
        return ConfigBasedTemplate(template_config)
    
    def get_available_templates(self) -> List[str]:
        """获取可用模板列表"""
        return list(self.config['templates'].keys())

class UniversalDocumentGenerator:
    """通用文档生成器"""
    
    def __init__(self, template_type: str, config_path: str = "templates_config.yaml"):
        self.logger = self._setup_logger()
        self.factory = TemplateFactory(config_path)
        self.template = self.factory.create_template(template_type)
        self.style_config = self.template.get_style_config()
        self.template_type = template_type
        
        self.logger.info(f"初始化模板: {template_type}")
        self.hyperlink_manager = DocumentHyperlinkManager()
    


    def generate_document(self, data: List[dict]) -> str:
        """生成Word文档"""
        try:
            self.logger.info(f"开始生成文档，数据条数: {len(data)}")
            
            # 1. 初始化文档
            doc = Document(self.template.get_template_path())
            written_heading_1 = set()
            
            # 2. 处理数据
            for i, item in enumerate(data):
                try:
                    # 使用模板特定的内容处理逻辑
                    processed_item = self.template.process_content(item)
                    
                    # 应用通用的文档生成逻辑
                    self._add_content_to_document(doc, processed_item, written_heading_1)
                    
                except Exception as e:
                    self.logger.warning(f"处理第 {i+1} 条数据时出错: {e}")
                    continue
            
            # 3. 后处理（仅文本格式化）
            self._post_process_document(doc)
            
            # 4. 保存文档
            doc_path = self._save_document(doc)
            
            # 5. 追加结尾模板
            if self.template.get_end_template_path():
                self._append_end_template(doc_path)

            # 6. 创建返回目录的超链接
            self.hyperlink_manager.create_return_to_toc_hyperlinks(
                doc_path, 
                self.style_config.get('target_bookmark', '目录'),
                self.style_config.get('return_link_text', '返回目录')
            )
            
            self.logger.info(f"文档生成成功: {doc_path}")
            return doc_path
            
        except Exception as e:
            self.logger.error(f"文档生成失败: {e}")
            raise
    
    def _add_content_to_document(self, doc, item: dict, written_heading_1: set):
        """添加内容到文档"""
        # 跳过完全空白的行
        if not any(item.get(key) for key in ['heading_1', 'heading_2', 'heading_3', 'title', 'content']):
            return
        
        # 处理一级标题
        heading_1 = item.get("heading_1")
        # 安全检查heading_1
        try:
            has_heading_1 = heading_1 and pd.notna(heading_1) and heading_1 not in written_heading_1
        except (ValueError, TypeError):
            has_heading_1 = heading_1 and heading_1 is not None and heading_1 not in written_heading_1
        
        if has_heading_1:
            self._add_heading(doc, heading_1, "heading_1")
            written_heading_1.add(heading_1)
        
        # 处理二级标题
        heading_2 = item.get("heading_2")
        try:
            has_heading_2 = heading_2 and pd.notna(heading_2)
        except (ValueError, TypeError):
            has_heading_2 = heading_2 and heading_2 is not None
        
        if has_heading_2:
            self._add_heading(doc, heading_2, "heading_2")
        
        # 处理三级标题 - 优先使用heading_3字段，如果为空则使用title字段
        heading_3 = item.get("heading_3")
        title = item.get("title")
        
        # 检查heading_3字段
        try:
            has_heading_3 = heading_3 and pd.notna(heading_3)
        except (ValueError, TypeError):
            has_heading_3 = heading_3 and heading_3 is not None
        
        # 检查title字段
        try:
            has_title = title and pd.notna(title)
        except (ValueError, TypeError):
            has_title = title and title is not None
        
        # 优先使用heading_3作为三级标题，如果没有则使用title
        if has_heading_3:
            self._add_heading(doc, heading_3, "heading_3")
            self.logger.info(f"添加三级标题(heading_3): {heading_3}")
        elif has_title:
            self._add_heading(doc, title, "title")
            self.logger.info(f"添加三级标题(title): {title}")
        
        # 处理正文内容
        content = item.get("content")
        try:
            has_content = content and pd.notna(content)
        except (ValueError, TypeError):
            has_content = content and content is not None
        
        if has_content:
            self._add_content(doc, content)
        
        # 处理source和date字段，独立于content处理
        source = item.get("source")
        date = item.get("date")
        
        # 检查source是否有效
        try:
            has_source = source and pd.notna(source) and str(source).strip()
        except (ValueError, TypeError):
            has_source = source and source is not None and str(source).strip()
        
        # 检查date是否有效
        try:
            has_date = date and pd.notna(date) and str(date).strip()
        except (ValueError, TypeError):
            has_date = date and date is not None and str(date).strip()
        
        # 如果同时有source和date，使用组合格式
        if has_source and has_date:
            self.logger.info(f"添加source和date: {source}, {date}")
            self._add_source_and_date(doc, source, date)
        elif has_source:
            self.logger.info(f"添加source: {source}")
            self._add_source_only(doc, source)
        elif has_date:
            self.logger.info(f"添加date: {date}")
            self._add_date_only(doc, date)
        
        # 变更点：修改占位符插入逻辑，确保每个有效文章（有内容的条目）后都有返回目录
        # 这样可以为所有包含实际内容的条目提供导航功能，提升用户体验
        if has_content:
            add_return_directory_placeholder(doc, self.style_config.get('return_link', {}))
    
    def _add_heading(self, doc, text: str, heading_type: str):
        """添加标题"""
        config = self.style_config.get(heading_type, {})
        # 使用通用工具格式化标题文本与级别
        mapping = self.style_config.get('title_mapping', {})
        level = compute_heading_level(heading_type, mapping)
        formatted_text = format_title_text(text, heading_type)
        if level is None:
            paragraph = doc.add_paragraph(formatted_text)
        else:
            paragraph = doc.add_heading(formatted_text, level=level)
        # 仅当配置明确指定时应用段落格式
        if config and any(key in config for key in ['alignment', 'line_spacing', 'space_after', 'first_line_indent']):
            apply_paragraph_format(paragraph, config)
        # 字体属性
        for run in paragraph.runs:
            if config.get("font_name"):
                run.font.name = config["font_name"]
            if config.get("font_size"):
                run.font.size = Pt(config["font_size"])
            if config.get("bold") is not None:
                run.font.bold = config["bold"]
        return paragraph
    
    def _is_likely_image_url(self, url: str) -> bool:
        """
        判断URL是否可能是图片链接
        基于文件扩展名和域名特征进行预筛选
        """
        if not url:
            return False
        
        # 图片URL特征模式（常见图片扩展名或图片服务域名）
        import re
        image_url_indicators = re.compile(r'(?i)\.(jpg|jpeg|png|gif|bmp|webp|svg)(?:[?#].*)?$|(?:img\.|image\.|pic\.|photo\.|cdn\.|oss\.|qpic\.)')
        
        # 检查是否包含图片扩展名或图片服务特征
        if image_url_indicators.search(url):
            return True
        
        # 排除明显的网页链接特征
        webpage_indicators = ['#/', '?page=', '/article/', '/news/', '/content/', '.html', '.htm', '.php', '.asp']
        for indicator in webpage_indicators:
            if indicator in url.lower():
                return False
        
        # 如果无明确特征，保守处理：不作为图片
        return False
    
    def _sanitize_url(self, raw: str):
        """
        清洗 URL 尾随标点，返回净化后的 URL 与被剥离的尾缀
        目的：在正文中 URL 常被括号/引号/句读符包裹，例如："(https://xx/a.png)"、"https://xx/a.png，"
        这些右侧标点若被包含入匹配结果，会导致下载失败或 404。
        参数：
            raw (str): 原始匹配到的 URL 片段（可能包含尾随标点）
        返回：
            Tuple[str, str]: (clean_url, trailing_suffix)
        关键逻辑：
            - 仅从右侧迭代剥离"分隔型"标点（半/全角），不影响查询参数字符（? & = % # 等）
            - 标点集合包含：), ], }, 半角/全角逗号句号问号冒号分号感叹号，引号，中文右括号等
        边界：若 URL 末尾本身为合法字符，将不被剥离；若全部为标点，保留原始串以避免意外吞没。
        """
        trailing = ""
        if not raw:
            return raw, trailing
        # 定义可能出现在右侧且不应参与下载的标点（含全角）
        trailing_set = set(list('),.;:!?\'\"]>}'))
        for ch in ['，', '。', '；', '：', '！', '？', '"', "'", '》', '】', '）', '、', '＞', '』', '」']:
            trailing_set.add(ch)
        s = raw
        # 迭代剥离右侧连续标点
        while s and s[-1] in trailing_set:
            trailing = s[-1] + trailing
            s = s[:-1]
        # 保底：若剥离后为空，回退为原串，避免将整段识别为空 URL
        if not s:
            return raw, ""
        return s, trailing
    
    def _add_content(self, doc, content):
        """添加正文内容，保持原有分段结构，并在拆段前内联处理URL为图片占位符，按原位置插入图片。
        说明：不再依赖 csv_to_word.parse_content 的“链接抽离再合并”策略，避免图片被统一追加到文末。
        同时在预处理阶段清除形如（688333.SH）的股票代码信息（兼容全/半角括号与中点）。
        """
        config = self.style_config.get("content", {})
        image_cfg = self.style_config.get("image", {})  # 可选图片配置

        # 1) 将 content 归一化为字符串，并在原文中就地进行 URL 替换为 [图片: 本地路径]
        import re
        from io import BytesIO
        import os
        try:
            import pandas as _pd
        except Exception:
            _pd = None
        # 处理 NaN/None
        if content is None:
            paragraphs_text = []
        else:
            if _pd is not None:
                try:
                    if _pd.isna(content):
                        paragraphs_text = []
                        content = ""
                    else:
                        content = str(content)
                except Exception:
                    content = str(content)
            else:
                content = str(content)

            # 预处理：清除股票代码信息（兼容全/半角括号与中点），如（688333.SH）/(3931.HK)
            # 同时清除股票价格信息，如(18.810, -0.09, -0.48%)
            content = re.sub(r'[（(]\s*\d{4,6}\s*[．\.]\s*[A-Za-z]{1,5}\s*[)）]', '', content)
            content = re.sub(r'[（(]\s*\d+\.\d+\s*,\s*[+-]?\d+\.\d+\s*,\s*[+-]?\d+\.\d+%\s*[)）]', '', content)

            # 改进的URL模式匹配：增加图片URL预筛选
            url_pattern = re.compile(r'https?://[^\s\u4e00-\u9fff]+')
            # 图片URL特征模式（常见图片扩展名或图片服务域名）
            image_url_indicators = re.compile(r'(?i)\.(jpg|jpeg|png|gif|bmp|webp|svg)(?:[?#].*)?$|(?:img\.|image\.|pic\.|photo\.|cdn\.|oss\.|qpic\.)')

            os.makedirs('temp-images', exist_ok=True)

            def _gen_filename(fmt: str) -> str:
                # 生成临时图片文件名（按格式扩展名），并返回绝对路径
                import time, random, os
                fname = f"{int(time.time())}_{random.randint(100000,999999)}.{fmt.lower()}"
                return os.path.abspath(os.path.join('temp-images', fname))



            def _replace_url(m: re.Match) -> str:
                """将URL就地替换为本地图片占位符；失败则保留原URL。
                - 增加URL预筛选，避免网页链接被误判为图片
                - 使用合理的请求头，提升下载成功率
                - 对不被 python-docx 原生支持的格式（如 WEBP/AVIF/TIFF 等）统一转为 PNG
                - 对含透明通道的图片优先保存为 PNG；JPEG 强制转为 RGB
                - 增加重试机制和更详细的错误处理
                """
                orig_url = m.group(0)
                # 先剥离尾随全/半角标点，避免下载失败，同时将尾缀在占位符后还原
                url, trailing_suffix = self._sanitize_url(orig_url)
                
                # 预筛选：检查URL是否可能是图片链接
                if not self._is_likely_image_url(url):
                    if hasattr(self, 'logger') and self.logger:
                        self.logger.info(f"跳过非图片URL: {url}")
                    return orig_url  # 保留原URL，不进行图片处理
                
                local_path = None
                max_retries = 5  # 增加重试次数从2次到5次
                
                for attempt in range(max_retries + 1):
                    try:
                        import requests
                        from PIL import Image
                        # 基于测试结果优化的请求头配置（80%成功率）
                        import random
                        
                        # 动态User-Agent池（基于测试验证的有效选项）
                        user_agents = [
                            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
                            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/121.0',
                            'Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1',
                            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                        ]
                        
                        headers = {
                            'User-Agent': random.choice(user_agents),
                            'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
                            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Cache-Control': 'no-cache',
                            'Pragma': 'no-cache',
                            'Sec-Fetch-Dest': 'image',
                            'Sec-Fetch-Mode': 'no-cors',
                            'Sec-Fetch-Site': 'cross-site'
                        }
                        
                        # 针对新浪图片的优化Referer策略（基于80%成功率测试）
                        if 'sinaimg.cn' in url.lower() or 'sina.com' in url.lower():
                            # 新浪图片专用Referer池
                            sina_referers = [
                                'https://finance.sina.com.cn/',
                                'https://news.sina.com.cn/',
                                'https://www.sina.com.cn/',
                                'https://mobile.sina.com.cn/',
                                'https://sina.com.cn/'
                            ]
                            headers['Referer'] = random.choice(sina_referers)
                        else:
                            # 对于其他域名，使用通用策略
                            from urllib.parse import urlparse
                            parsed = urlparse(url)
                            headers['Referer'] = f"{parsed.scheme}://{parsed.netloc}/"
                        
                        # 增加超时时间和重试逻辑
                        timeout = 20 if attempt == 0 else 30 + (attempt * 10)  # 递增超时时间
                        resp = requests.get(url, timeout=timeout, headers=headers, stream=True)
                        resp.raise_for_status()
                        
                        # 优化Content-Type检查逻辑
                        content_type = (resp.headers.get('Content-Type') or '').lower()
                        if 'image' not in content_type:
                            # 严格检查：如果Content-Type不包含image，直接跳过
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"非图片内容类型: {content_type} for {url}, 跳过处理")
                            return orig_url  # 返回原URL，不进行图片处理
                        
                        # 检查是否为HTML内容（常见的误判情况）
                        if any(html_type in content_type for html_type in ['text/html', 'text/plain', 'application/json']):
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"检测到网页内容: {content_type} for {url}, 跳过处理")
                            return orig_url
                        
                        data = resp.content
                        if len(data) < 100:  # 检查文件大小（100字节阈值属于经验值）
                            raise ValueError(f"图片文件过小: {len(data)} bytes")
                        
                        img = Image.open(BytesIO(data))
                        
                        # 检查图片尺寸
                        if img.size[0] < 10 or img.size[1] < 10:
                            raise ValueError(f"图片尺寸过小: {img.size}")
                        
                        fmt = (img.format or '').upper()
                        has_alpha = 'A' in img.getbands()
                        
                        # 选择保存格式
                        if fmt in ('WEBP', 'AVIF', 'TIFF', 'BMP', 'ICO') or has_alpha:
                            save_fmt = 'PNG'
                        else:
                            save_fmt = fmt if fmt in ('JPEG', 'PNG') else 'JPEG'
                        
                        # 模式转换
                        if save_fmt == 'JPEG' and img.mode != 'RGB':
                            img = img.convert('RGB')
                        elif save_fmt == 'PNG' and img.mode not in ('RGBA', 'RGB'):
                            img = img.convert('RGBA' if has_alpha else 'RGB')
                        
                        # 确保目录存在
                        os.makedirs('temp-images', exist_ok=True)
                        local_path = _gen_filename(save_fmt)
                        img.save(local_path, format=save_fmt, quality=95 if save_fmt == 'JPEG' else None)
                        
                        if hasattr(self, 'logger') and self.logger:
                            self.logger.info(f"下载并内联图片: {url} -> {local_path} (fmt={save_fmt}, size={img.size})")
                        return f"[图片: {local_path}]" + (trailing_suffix or '')
                        
                    except requests.exceptions.Timeout as e:
                        if attempt < max_retries:
                            wait_time = 2 ** attempt  # 指数退避
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"图片下载超时，等待{wait_time}秒后重试 {attempt + 1}/{max_retries}: {url}")
                            time.sleep(wait_time)
                            continue
                        else:
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"图片下载超时，已达最大重试次数: {url}")
                            return orig_url
                    except requests.exceptions.RequestException as e:
                        if attempt < max_retries:
                            wait_time = 2 ** attempt  # 指数退避
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"图片下载网络错误，等待{wait_time}秒后重试 {attempt + 1}/{max_retries}: {url}, 错误: {e}")
                            time.sleep(wait_time)
                            continue
                        else:
                            if hasattr(self, 'logger') and self.logger:
                                self.logger.warning(f"图片下载网络错误，已达最大重试次数: {url}, 错误: {e}")
                            return orig_url
                    except Exception as e:
                        if hasattr(self, 'logger') and self.logger:
                            self.logger.warning(f"图片下载或转换失败，保留原链接: {url}, 错误: {e}")
                        return orig_url
                
                # 如果所有重试都失败，返回原URL
                return orig_url
            text = content
            try:
                text = url_pattern.sub(_replace_url, text)
            except Exception as e:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.error(f"URL 替换阶段出错: {e}")
            # 轻量清洗：去除全角空格，规整空白与换行，避免出现大量“　　”以及异常空行
            try:
                # 转换"^l"为"^p"（Word换行符标准化）
                text = text.replace('^l', '^p')
                # 移除全角空格（U+3000）
                text = text.replace('　', '')
                # 合并连续半角空格/制表符为单个空格
                text = re.sub(r'[ \t]+', ' ', text)
                # 行内首尾空白去除
                text = '\n'.join(line.strip() for line in text.split('\n'))
                # 压缩三个及以上连续空行为两个空行（用于段落分段）
                text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            except Exception as e:
                if hasattr(self, 'logger') and self.logger:
                    self.logger.warning(f"文本清洗阶段出错: {e}")
            # 2) 轻量分段：按两个及以上换行分段；若无则整体作为一个段落
            parts = [p.strip() for p in re.split(r"\n{2,}", text) if p and p.strip()]
            paragraphs_text = parts if parts else ([text.strip()] if text.strip() else [])

        for para_text in paragraphs_text:
            if not para_text.strip():
                continue

            # 如果段落包含图片占位符或是单独的本地图片行，则按顺序插入文本与图片
            if self._is_image_marker(para_text) or para_text.strip().startswith("temp-images"):
                marker_text = para_text
                if para_text.strip().startswith("temp-images"):
                    # 兼容旧格式：纯图片路径行 -> 统一转为占位符格式
                    marker_text = f"[图片: {para_text.strip()}]"
                self._add_text_and_images_by_marker(doc, marker_text, config, image_cfg)
                continue

            # 否则作为纯文本段落处理
            # 若包含单换行，则将每一行作为独立段落写入，避免软换行导致的“换行缺失”观感
            if "\n" in para_text:
                for line in (l for l in para_text.split('\n') if l.strip()):
                    paragraph = doc.add_paragraph()
                    paragraph.style = "Normal"
                    if config and any(key in config for key in ['alignment', 'line_spacing', 'space_after', 'first_line_indent']):
                        apply_paragraph_format(paragraph, config)
                    if "*" in line:
                        self._add_formatted_text(paragraph, line)
                    else:
                        run = paragraph.add_run(line)
                        # 启发式：疑似“内联段落标题”则自动加粗
                        try:
                            if self._looks_like_inline_heading(line):
                                run.bold = True
                        except Exception:
                            pass
                        if config and config.get("font_name"):
                            run.font.name = config["font_name"]
                        if config and config.get("font_size"):
                            run.font.size = Pt(config["font_size"]) 
            else:
                paragraph = doc.add_paragraph()
                paragraph.style = "Normal"
                if config and any(key in config for key in ['alignment', 'line_spacing', 'space_after', 'first_line_indent']):
                    apply_paragraph_format(paragraph, config)
                if "*" in para_text:
                    self._add_formatted_text(paragraph, para_text)
                else:
                    run = paragraph.add_run(para_text)
                    # 启发式：疑似“内联段落标题”则自动加粗
                    try:
                        if self._looks_like_inline_heading(para_text):
                            run.bold = True
                    except Exception:
                        pass
                    if config and config.get("font_name"):
                        run.font.name = config["font_name"]
                    if config and config.get("font_size"):
                        run.font.size = Pt(config["font_size"])

    def _is_image_marker(self, text: str) -> bool:
        """是否包含图片占位符，如：[图片: temp-images/xxx.jpg] 或 [图片: http://...]"""
        import re
        return bool(re.search(r'\[图片:\s*[^\]]+\]', text.strip()))

    def _add_text_and_images_by_marker(self, doc, marker_text: str, text_cfg: dict, image_cfg: dict):
        """将含有图片占位符的文本拆分为文本与图片，按出现顺序依次插入。
        - 文本段落：沿用正文样式与加粗格式处理
        - 图片段落：独立成段、支持对齐与最大宽高配置
        """
        import re, os
        from PIL import Image

        # 拆分出文本与图片路径（使用捕获组保持图片路径）
        pattern = r'\[图片:\s*([^\]]+)\]'
        parts = re.split(pattern, marker_text)

        i = 0
        while i < len(parts):
            # 文本部分
            text_part = parts[i]
            if text_part and text_part.strip():
                paragraph = doc.add_paragraph()
                paragraph.style = "Normal"
                if text_cfg and any(key in text_cfg for key in ['alignment', 'line_spacing', 'space_after', 'first_line_indent']):
                    apply_paragraph_format(paragraph, text_cfg)
                # 支持加粗标记与换行
                if "*" in text_part:
                    self._add_formatted_text(paragraph, text_part)
                else:
                    lines = text_part.split('\n')
                    for j, line in enumerate(lines):
                        if not line.strip():
                            continue
                        if j > 0:
                            paragraph.add_run().add_break()
                        run = paragraph.add_run(line)
                        if text_cfg.get("font_name"):
                            run.font.name = text_cfg["font_name"]
                    if text_cfg.get("font_size"):
                        run.font.size = Pt(text_cfg["font_size"])

            # 图片部分（parts[i+1] 为路径）
            if i + 1 < len(parts):
                image_path = (parts[i + 1] or "").strip()
                if image_path:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    try:
                        if not os.path.exists(image_path):
                            # 尝试兼容相对路径位于 temp-images 目录的情况
                            alt_path = os.path.join(os.getcwd(), image_path)
                            if os.path.exists(alt_path):
                                image_path = alt_path
                        # 计算显示尺寸
                        with Image.open(image_path) as img:
                            width_px, height_px = img.size
                        max_w_in = float(image_cfg.get('max_width_inches', 5.0))
                        max_h_in = float(image_cfg.get('max_height_inches', 4.0))
                        dpi = 96.0
                        img_w_in = width_px / dpi
                        img_h_in = height_px / dpi
                        aspect = img_w_in / img_h_in if img_h_in else 1.0
                        # 先按宽度约束
                        disp_w_in = min(max_w_in, img_w_in)
                        disp_h_in = disp_w_in / aspect if aspect else img_h_in
                        # 再按高度约束
                        if disp_h_in > max_h_in:
                            disp_h_in = max_h_in
                            disp_w_in = disp_h_in * aspect
                        run.add_picture(image_path, width=Inches(disp_w_in), height=Inches(disp_h_in))
                        # 对齐设置：居中对齐且首行不缩进
                        align = (image_cfg.get('alignment') or 'center').lower()
                        if align == 'left':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif align == 'right':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 设置首行不缩进（图片段落专用格式）
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.first_line_indent = Inches(0)  # 首行不缩进
                        if hasattr(self, 'logger') and self.logger:
                            self.logger.info(f"插入图片: {image_path} -> {disp_w_in:.2f}x{disp_h_in:.2f} in")
                    except Exception as e:
                        if hasattr(self, 'logger') and self.logger:
                            self.logger.error(f"插入图片失败 {image_path}: {e}")
                        paragraph.add_run(f"[图片加载失败: {os.path.basename(image_path) if image_path else '未知图片'}]")
            i += 2

    def _add_formatted_text(self, paragraph, text: str):
        """添加带格式的文本"""
        import re
        text_chunks = re.split(r"(\*\*.*?\*\*)", text)
        
        for chunk in text_chunks:
            if chunk.startswith("**") and chunk.endswith("**"):
                # 粗体文本
                chunk = chunk[2:-2]
                run = paragraph.add_run(chunk)
                run.bold = True
            else:
                # 普通文本
                chunk = chunk.replace("*", "")
                run = paragraph.add_run(chunk)
            
            # 只有在配置中明确指定字体属性时才设置字体
            config = self.style_config.get("content", {})
            if config.get("font_name"):
                run.font.name = config["font_name"]
            if config.get("font_size"):
                run.font.size = Pt(config["font_size"]) 
    
    def _looks_like_inline_heading(self, text: str) -> bool:
        """启发式判断文本是否为"内联段落标题"。
        规则说明：
        - 长度区间 [6, 50]（扩大上限以支持更长的标题）
        - 不以句号/问号/叹号等断句标点结尾（中英文）
        - 行内不应包含明显断句标点（但允许冒号）
        - 不以 URL 前缀开头
        - 至少包含 2 个中文字符
        - 标点占比不过高
        - 特殊模式：支持"关键词：描述性内容"格式（如"新房：日均增长7％，改善客群向外环突围"）
        返回 True 表示更像标题，应加粗。
        """
        try:
            t = str(text).strip()
            if len(t) < 6 or len(t) > 50:  # 扩大长度上限
                return False
            lower = t.lower()
            if lower.startswith("http://") or lower.startswith("https://"):
                return False
            
            # 特殊模式：检查是否为"关键词：描述"格式
            if "：" in t or ":" in t:
                # 分割冒号前后部分
                colon_parts = t.split("：") if "：" in t else t.split(":")
                if len(colon_parts) == 2:
                    prefix, suffix = colon_parts[0].strip(), colon_parts[1].strip()
                    # 前缀应该是简短的关键词（2-8个字符）
                    if 2 <= len(prefix) <= 8:
                        # 后缀可以包含逗号等标点，但不应以句号结尾
                        if not suffix.endswith(tuple("。；;.!！？?")):
                            # 检查中文字符数量
                            chinese_count = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
                            if chinese_count >= 3:  # 至少3个中文字符
                                return True
            
            # 原有规则：末尾不应为断句标点
            if t.endswith(tuple("。；;.!！？?:：")):
                return False
            # 行内不应包含明显断句标点（但允许冒号）
            for ch in "，、；。.!！？":
                if ch in t:
                    return False
            chinese_count = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
            if chinese_count < 2:
                return False
            puncts = "''()（）[]【】<>《》—-··+/:\""
            punct_count = sum(1 for ch in t if ch in puncts)
            if punct_count > max(2, len(t) // 6):
                return False
            return True
        except Exception:
            return False
    
    def _add_source_and_date(self, doc, source: str, date: str):
        """添加信源和日期"""
        # 优先使用source_date样式配置，如果没有则使用content配置
        config = self.style_config.get("source_date", self.style_config.get("content", {}))
        paragraph = doc.add_paragraph()
        
        source_date_text = f"{source} {date}"
        run = paragraph.add_run(source_date_text)
        
        # 设置字体属性
        run.font.name = config.get("font_name", "宋体")
        run.font.size = Pt(config.get("font_size", 12))
        
        # 应用段落格式
        if config:
            apply_paragraph_format(paragraph, config)
        else:
            paragraph.paragraph_format.space_after = Pt(12)
    
    def _add_date(self, doc, date: str):
        """添加日期"""
        config = self.style_config.get("content", {})
        paragraph = doc.add_paragraph()
        paragraph.style = "Normal"
        
        run = paragraph.add_run(date)
        run.font.name = config.get("font_name", "宋体")
        run.font.size = Pt(config.get("font_size", 12))
        
        self._apply_paragraph_format(paragraph, config)
        paragraph.paragraph_format.space_after = Pt(12)
    
    def _add_date_only(self, doc, date: str):
        """只添加日期(当source为空时)"""
        paragraph = doc.add_paragraph()
        paragraph.style = "Normal"
        
        # 添加日期文本
        run = paragraph.add_run(str(date).strip())
        
        # 设置字体为Times New Roman小四
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)  # 小四号字体
        
        # 设置段落格式：首行缩进、两端对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
        paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        paragraph.paragraph_format.space_after = Pt(12)  # 段后间距
    
    def _add_return_directory_placeholder(self, doc):
        """添加返回目录占位符"""
        try:
            template_config = self.template.get_style_config()
            return_link_config = template_config.get('return_link', {})
            
            # 获取配置参数
            return_text = return_link_config.get('text', '返回目录')
            font_name = return_link_config.get('font_name', '宋体')
            font_size = return_link_config.get('font_size', 12)
            alignment = return_link_config.get('alignment', 'right')
            
            # 创建占位符段落
            return_paragraph = doc.add_paragraph()
            return_paragraph.style = "Normal"
            
            # 设置对齐方式
            if alignment == 'right':
                return_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'center':
                return_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                return_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            return_paragraph.paragraph_format.line_spacing = 1.0
            return_paragraph.paragraph_format.space_after = Pt(12)
            
            # 添加占位符文本
            return_run = return_paragraph.add_run(return_text)
            return_run.font.name = font_name
            return_run.font.size = Pt(font_size)
            
            # 如果配置了下划线，添加下划线
            if return_link_config.get('underline', False):
                return_run.font.underline = WD_UNDERLINE.SINGLE
                
            self.logger.info(f"已添加返回目录占位符: '{return_text}'")
            
        except Exception as e:
            self.logger.error(f"添加返回目录占位符失败: {e}")

    def _add_return_to_directory_simple(self, doc):
        """添加简单的返回目录占位符（包装 utils.add_return_directory_placeholder）"""
        try:
            add_return_directory_placeholder(doc, {
                'text': '返回目录',
                'font_name': '宋体',
                'font_size': 12,
                'alignment': 'right',
                'underline': True,
            })
            self.logger.info("已添加返回目录占位符（简单实现包装）")
        except Exception as e:
            self.logger.error(f"添加返回目录链接失败: {e}")

    def _add_return_link(self, doc):
        """
        添加返回目录文本（包装 utils.add_return_directory_placeholder）
        根据模板配置动态生成文本内容和格式；在后处理阶段会统一为这些文本添加超链接
        """
        try:
            template_config = self.template.get_style_config()
            return_link_config = template_config.get('return_link', {})
            add_return_directory_placeholder(doc, return_link_config)
            self.logger.info("已根据模板配置添加返回目录占位符")
        except Exception as e:
            self.logger.error(f"添加返回目录文本失败: {e}")

    def _ensure_bookmark_exists(self, doc, bookmark_name: str):
        """确保指定的书签存在,如果不存在则查找关键词并设置书签"""
        from docx.oxml.shared import qn
        from docx.oxml import parse_xml
        
        # 检查书签是否已存在
        existing_bookmarks = []
        for element in doc.element.iter():
            if element.tag.endswith('bookmarkStart'):
                name_attr = element.get(qn('w:name'))
                if name_attr:
                    existing_bookmarks.append(name_attr)
        
        if bookmark_name in existing_bookmarks:
            self.logger.info(f"书签 '{bookmark_name}' 已存在")
            return
        
        # 遍历文档查找目标关键词
        for paragraph in doc.paragraphs:
            if bookmark_name in paragraph.text:
                self.logger.info(f"找到关键词 '{bookmark_name}',正在设置书签")
                add_bookmark_to_paragraph_xml(paragraph, bookmark_name)
                return
        
        # 如果在正文中没找到,检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if bookmark_name in paragraph.text:
                            self.logger.info(f"在表格中找到关键词 '{bookmark_name}',正在设置书签")
                            add_bookmark_to_paragraph_xml(paragraph, bookmark_name)
                            return
        
        self.logger.warning(f"未找到关键词 '{bookmark_name}',无法设置书签")
    
    def _add_bookmark_to_paragraph(self, paragraph, bookmark_name: str):
        """在段落中添加书签（包装 utils 实现，保持原格式不变）"""
        add_bookmark_to_paragraph_xml(paragraph, bookmark_name)
        self.logger.info(f"成功为段落设置书签: {bookmark_name},保持原有格式")
    
    def _apply_paragraph_format(self, paragraph, config: dict):
        """应用段落格式（包装 utils.apply_paragraph_format）"""
        return apply_paragraph_format(paragraph, config)
    
    def _post_process_document(self, doc):
        """文档后处理，包含完整的返回目录功能实现"""
        try:
            # 应用文本处理函数

            center_image_description_paragraphs(doc)
            # 关键修复：实现完整的返回目录功能
            # 读取配置
            template_config = self.template.get_style_config()
            target_bookmark = template_config.get('target_bookmark', '目录')
            return_link_config = template_config.get('return_link', {})
            return_text = return_link_config.get('text', '返回目录')
            
            self.logger.info(f"开始处理“返回目录”超链接，目标书签: '{target_bookmark}'")
            # 不再动态创建书签，直接使用模板中预设的书签
            # 调用重构后的辅助函数，该函数现在只负责创建链接
            convert_return_placeholders_to_hyperlinks(doc, return_text, target_bookmark)
            
            self.logger.info("文档后处理完成")
        except Exception as e:
            self.logger.warning(f"文档后处理出错: {e}")

    def _implement_return_directory_functionality(self, doc):
        """实现完整的返回目录功能"""
        try:
            # 获取模板配置
            template_config = self.template.get_style_config()
            target_bookmark = template_config.get('target_bookmark', '目录')
            return_link_config = template_config.get('return_link', {})
            return_text = return_link_config.get('text', '返回目录')
            
            self.logger.info(f"开始处理返回目录功能，目标书签: '{target_bookmark}', 返回文本: '{return_text}'")
            
            # 第一步：查找目标关键词并创建书签
            bookmark_created = self._create_target_bookmark(doc, target_bookmark)
            
            if not bookmark_created:
                self.logger.warning(f"未找到目标关键词 '{target_bookmark}'，使用默认书签")
                target_bookmark = "default_toc"  # 使用默认书签名
            
            # 第二步：将所有占位符转换为超链接
            converted_count = self._convert_placeholders_to_hyperlinks(doc, return_text, target_bookmark, return_link_config)
            
            self.logger.info(f"返回目录功能处理完成，转换了 {converted_count} 个占位符为超链接")
            
        except Exception as e:
            self.logger.error(f"返回目录功能实现失败: {e}")

    def _create_target_bookmark(self, doc, target_text):
        """包装 utils.create_target_bookmark_by_keyword：在包含目标关键词的段落插入书签。"""
        try:
            return create_target_bookmark_by_keyword(doc, target_text)
        except Exception as e:
            self.logger.error(f"创建目标书签失败: {e}")
            return False

    def _convert_placeholders_to_hyperlinks(self, doc, return_text, target_bookmark, return_link_config):
        """包装 utils.convert_return_placeholders_to_hyperlinks：将占位符转换为指向书签的内部超链接。"""
        try:
            return convert_return_placeholders_to_hyperlinks(doc, return_text, target_bookmark, return_link_config)
        except Exception as e:
            self.logger.error(f"转换返回目录占位符失败: {e}")
            return 0

    def _save_document(self, doc) -> str:
        """保存文档"""
        os.makedirs("temp-data", exist_ok=True)
        timestamp = int(time.time())
        random_suffix = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        doc_path = f"temp-data/{timestamp}_{random_suffix}_{self.template_type}.docx"
        doc.save(doc_path)
        return doc_path
    
    def _append_end_template(self, doc_path: str):
        """
        使用 docxcompose 库追加结尾模板，以保留其完整的格式和页面布局。

        Args:
            doc_path (str): 主文档的路径。
        """
        try:
            master_doc = Document(doc_path)
            composer = Composer(master_doc)
            
            end_template_path = self.template.get_end_template_path()

            if not os.path.exists(end_template_path):
                self.logger.warning(f"结尾模板文件不存在: {end_template_path}")
                return

            end_template_doc = Document(end_template_path)
            composer.append(end_template_doc)
            
            composer.save(doc_path)
            self.logger.info(f"已使用 docxcompose 成功将结尾模板追加到: {doc_path}")

        except Exception as e:
            self.logger.error(f"使用 docxcompose 追加结尾模板时发生错误: {e}")
            # 作为备选方案，可以回退到旧的、仅复制内容的方法，或者直接抛出异常
            # 这里选择记录错误并继续，以避免完全失败
            pass
    
    def _setup_logger(self) -> logging.Logger:
        """设置日志器"""
        logger = logging.getLogger(f"UniversalDocumentGenerator")
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger

def csv_to_word_universal(csv_file: str, template_type: str, 
                         config_path: str = "templates_config.yaml") -> str:
    """通用CSV到Word转换函数
    
    Args:
        csv_file: CSV文件路径
        template_type: 模板类型
        config_path: 配置文件路径
    
    Returns:
        str: 生成的Word文档路径
    """
    # 读取CSV数据，尝试不同编码
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
    df = None
    
    for encoding in encodings:
        try:
            # 检查是否需要跳过第一行（如果第一行是"Table 1"等表头）
            first_line = None
            with open(csv_file, 'r', encoding=encoding) as f:
                first_line = f.readline().strip()
            
            # 如果第一行包含"Table"等关键词，跳过第一行
            skip_rows = 1 if first_line and ('Table' in first_line or 'table' in first_line) else 0
            
            df = pd.read_csv(csv_file, encoding=encoding, skiprows=skip_rows)
            print(f"成功使用编码 {encoding} 读取CSV文件，跳过行数: {skip_rows}")
            break
        except UnicodeDecodeError:
            continue
    
    if df is None:
        raise ValueError(f"无法读取CSV文件 {csv_file}，尝试了编码: {encodings}")
    
    # 应用字段映射，将中文字段名映射为英文字段名
    df = apply_csv_field_mapping(df)
    
    data = df.to_dict('records')
    
    # 创建文档生成器
    generator = UniversalDocumentGenerator(template_type, config_path)
    
    # 生成文档
    doc_path = generator.generate_document(data)
    
    return doc_path

if __name__ == "__main__":
    # 示例用法
    try:
        # 创建模板工厂
        factory = TemplateFactory()
        # 显示可用模板
        print("可用模板:", factory.get_available_templates())
        
        # 处理用户指定的国资委CSV文件
        print("\n=== 生成国资委文档 ===")
        csv_file = "temp-data/电力专项排版.csv"
        doc_path = csv_to_word_universal(csv_file, "electricity")
        print(f"国资委模板文档生成完成: {doc_path}")

    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()