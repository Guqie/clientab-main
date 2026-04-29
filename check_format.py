# -*- coding: utf-8 -*-
import sys
import re
from docx import Document

doc = Document(r'd:\桌面\战新与未来产业月报第四期0427.docx')

# 检查格式问题
issues = []

# 需要删除的词
delete_words = ['专家', '认为', '分析', '表示', '介绍']

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # 检查是否有需要删除的词
    for word in delete_words:
        if word in text:
            issues.append(f'行{i+1}: 包含"{word}" - {text[:100]}')
    
    # 检查标题格式（数字序号后应有空格）
    if re.match(r'^\d+\.', text):
        if len(text) > 2 and text[1] == '.' and text[2] != ' ':
            issues.append(f'行{i+1}: 标题序号后未加空格 - {text[:80]}')

# 检查中车相关
zhongche_words = ['中国中车', 'CRRC', '中车集团']
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for word in zhongche_words:
        if word in text:
            issues.append(f'行{i+1}: 疑似中车新闻 - {text[:100]}')

# 写入文件
with open('format_check_result.txt', 'w', encoding='utf-8') as f:
    f.write('=== 格式问题检查结果 ===\n')
    f.write(f'共发现 {len(issues)} 处问题\n\n')
    for issue in issues:
        f.write(issue + '\n')

print(f'检查完成，共发现 {len(issues)} 处问题，已保存到 format_check_result.txt')
