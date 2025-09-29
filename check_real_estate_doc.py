#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
房地产周刊文档检查脚本
用于验证图片格式设置和^l转^p转换功能
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_document():
    """
    检查生成的房地产周刊Word文档
    验证图片格式设置和字符转换功能
    """
    doc_path = 'temp-data/1758619325_386932_guoziwei.docx'
    print(f'检查文档: {doc_path}')
    
    if not os.path.exists(doc_path):
        print('❌ 文档不存在!')
        return False
    
    try:
        # 加载文档
        doc = Document(doc_path)
        print(f'✅ 文档加载成功，总段落数: {len(doc.paragraphs)}')
        
        # 初始化统计变量
        image_count = 0
        image_paragraphs = 0
        center_aligned_images = 0
        zero_indent_images = 0
        l_char_count = 0
        p_char_count = 0
        
        # 遍历所有段落
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            
            # 统计^l和^p字符
            l_char_count += text.count('^l')
            p_char_count += text.count('^p')
            
            # 检查是否包含图片
            has_image = False
            for run in paragraph.runs:
                # 检查是否有图片元素
                if run.element.xpath('.//a:blip'):
                    has_image = True
                    image_count += 1
                    break
            
            if has_image:
                image_paragraphs += 1
                
                # 检查段落对齐方式
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    center_aligned_images += 1
                
                # 检查首行缩进
                first_line_indent = paragraph.paragraph_format.first_line_indent
                if first_line_indent is None or first_line_indent == Inches(0):
                    zero_indent_images += 1
                
                print(f'  图片段落 {i+1}: 对齐={paragraph.alignment}, 缩进={first_line_indent}')
        
        # 输出统计结果
        print(f'\n=== 统计结果 ===')
        print(f'图片数量: {image_count}')
        print(f'包含图片的段落数: {image_paragraphs}')
        print(f'居中对齐的图片段落: {center_aligned_images}')
        print(f'首行缩进为0的图片段落: {zero_indent_images}')
        print(f'^l字符数量: {l_char_count}')
        print(f'^p字符数量: {p_char_count}')
        
        # 功能验证
        print(f'\n=== 功能验证 ===')
        success = True
        
        if image_paragraphs > 0:
            if center_aligned_images == image_paragraphs:
                print('✅ 图片居中对齐设置正确')
            else:
                print(f'❌ 图片居中对齐设置有问题: {center_aligned_images}/{image_paragraphs}')
                success = False
                
            if zero_indent_images == image_paragraphs:
                print('✅ 图片首行缩进设置正确')
            else:
                print(f'❌ 图片首行缩进设置有问题: {zero_indent_images}/{image_paragraphs}')
                success = False
        else:
            print('⚠️  文档中没有检测到图片')
        
        if l_char_count == 0:
            print('✅ ^l字符已正确转换')
        else:
            print(f'❌ 仍有{l_char_count}个^l字符未转换')
            success = False
        
        print(f'\n=== 测试总结 ===')
        if success:
            print('🎉 所有功能测试通过!')
        else:
            print('⚠️  部分功能需要检查')
            
        return success
        
    except Exception as e:
        print(f'❌ 检查文档时出错: {e}')
        return False

if __name__ == '__main__':
    check_document()