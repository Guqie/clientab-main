#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查段落0的详细信息
用于分析对齐方式为None的原因
"""

from docx import Document

def check_paragraph_zero():
    """
    检查段落0的详细信息
    """
    doc_path = 'temp-data/1758620917_204430_realty.docx'
    doc = Document(doc_path)
    
    print("段落0详细信息:")
    print("=" * 50)
    
    p = doc.paragraphs[0]
    print(f'文本: "{p.text}"')
    print(f'对齐方式: {p.alignment}')
    print(f'首行缩进: {p.paragraph_format.first_line_indent}')
    print(f'段落样式: {p.style.name if p.style else "None"}')
    print(f'段落长度: {len(p.text)}')
    print(f'runs数量: {len(p.runs)}')
    
    print("\nRuns详细信息:")
    for i, run in enumerate(p.runs):
        has_drawing = bool(run._element.xpath('.//w:drawing'))
        has_blip = bool(run._element.xpath('.//a:blip'))
        print(f'  run{i}: "{run.text}" (图片drawing: {has_drawing}, 图片blip: {has_blip})')
    
    # 检查前几个段落的情况
    print("\n前5个段落概览:")
    for i in range(min(5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        has_image = False
        for run in p.runs:
            if run._element.xpath('.//w:drawing') or run._element.xpath('.//a:blip'):
                has_image = True
                break
        print(f'段落{i}: 文本="{p.text[:20]}..." 对齐={p.alignment} 图片={has_image}')

if __name__ == "__main__":
    check_paragraph_zero()