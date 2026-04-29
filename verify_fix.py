# -*- coding: utf-8 -*-
"""
验证修复结果
"""
from docx import Document
import re

def verify():
    doc = Document(r'd:\桌面\战新与未来产业月报第四期0427_已修复.docx')
    
    print("=" * 60)
    print("验证修复结果")
    print("=" * 60)
    
    issues = []
    
    # 1. 检查标题格式
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^\d+\.', text):
            if text[1] == '.' and text[2] != ' ':
                issues.append(f"[标题] 行{i+1}: {text[:50]}")
    
    # 2. 检查关键词
    check_words = ['表示', '介绍，', '偏差分析', '水质分析实验室']
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for word in check_words:
            if word in text:
                issues.append(f"[关键词] 行{i+1}: 含'{word}'")
    
    # 3. 检查"认为"（除了免责声明）
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '认为' in text and '我们认为可信' not in text and '被认为' not in text:
            issues.append(f"[关键词] 行{i+1}: 含'认为'")
    
    # 4. 检查"被认为"
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '被认为' in text:
            issues.append(f"[关键词] 行{i+1}: 含'被认为'")
    
    if issues:
        print(f"\n发现 {len(issues)} 处问题:\n")
        for issue in issues[:20]:
            print(f"  {issue}")
    else:
        print("\n[PASS] 所有检查项通过!")
    
    # 统计段落
    total = len([p for p in doc.paragraphs if p.text.strip()])
    print(f"\n文档段落数: {total}")
    
    return len(issues) == 0

if __name__ == '__main__':
    verify()
